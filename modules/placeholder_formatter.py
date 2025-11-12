"""
Module pour gérer le remplacement de placeholders en préservant le formatage.
"""

import re
from typing import Dict, List, Tuple, Optional
from docx.shared import RGBColor


def find_placeholder_in_runs(runs: list, placeholder: str) -> Optional[Tuple[int, int, Dict]]:
    """
    Trouve un placeholder dans une liste de runs et retourne son formatage.

    Args:
        runs: Liste de runs du paragraphe
        placeholder: Le placeholder à chercher (avec les crochets)

    Returns:
        Tuple (start_run_idx, end_run_idx, format_dict) ou None si non trouvé
    """
    # Reconstituer le texte depuis les runs
    full_text = "".join(run.text for run in runs)

    if placeholder not in full_text:
        return None

    # Position du placeholder dans le texte complet
    placeholder_start = full_text.index(placeholder)
    placeholder_end = placeholder_start + len(placeholder)

    # Trouver les runs concernés
    current_pos = 0
    start_run_idx = None
    end_run_idx = None
    format_dict = {}

    for idx, run in enumerate(runs):
        run_len = len(run.text)
        run_end = current_pos + run_len

        # Ce run contient le début du placeholder?
        if start_run_idx is None and current_pos <= placeholder_start < run_end:
            start_run_idx = idx
            # Capturer le formatage du run où commence le placeholder
            format_dict = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            }

        # Ce run contient la fin du placeholder?
        if start_run_idx is not None and current_pos < placeholder_end <= run_end:
            end_run_idx = idx
            break

        current_pos = run_end

    if start_run_idx is not None and end_run_idx is not None:
        return (start_run_idx, end_run_idx, format_dict)

    return None


def replace_placeholder_with_format(paragraph, placeholder: str, replacement: str) -> bool:
    """
    Remplace un placeholder dans un paragraphe en préservant son formatage.

    Args:
        paragraph: Paragraphe docx
        placeholder: Le placeholder à remplacer (avec les crochets)
        replacement: Le texte de remplacement

    Returns:
        True si le remplacement a été effectué, False sinon
    """
    # Trouver le placeholder et son formatage
    result = find_placeholder_in_runs(paragraph.runs, placeholder)

    if not result:
        return False

    start_run_idx, end_run_idx, format_dict = result

    # Stratégie: On va reconstruire le paragraphe
    # 1. Construire le nouveau texte complet
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text.replace(placeholder, replacement)

    # 2. Vider tous les runs
    for run in paragraph.runs:
        run.text = ""

    # 3. Reconstruire en préservant le formatage
    # Pour simplifier, on va créer 3 segments:
    # - Texte avant le placeholder (formatage original)
    # - Remplacement (avec le formatage du placeholder)
    # - Texte après le placeholder (formatage original)

    placeholder_pos = full_text.index(placeholder)
    before_text = full_text[:placeholder_pos]
    after_text = full_text[placeholder_pos + len(placeholder):]

    # Texte avant
    if before_text:
        # Utiliser le formatage du premier run
        if paragraph.runs:
            run = paragraph.runs[0]
            run.text = before_text
        else:
            run = paragraph.add_run(before_text)

    # Remplacement avec formatage du placeholder
    run = paragraph.add_run(replacement)
    if format_dict['bold'] is not None:
        run.bold = format_dict['bold']
    if format_dict['italic'] is not None:
        run.italic = format_dict['italic']
    if format_dict['underline'] is not None:
        run.underline = format_dict['underline']
    if format_dict['font_name']:
        run.font.name = format_dict['font_name']
    if format_dict['font_size']:
        run.font.size = format_dict['font_size']
    if format_dict['font_color']:
        run.font.color.rgb = format_dict['font_color']

    # Texte après
    if after_text:
        run = paragraph.add_run(after_text)
        # Utiliser le formatage du dernier run original
        if len(paragraph.runs) > end_run_idx + 1:
            last_run = paragraph.runs[end_run_idx + 1]
            if last_run.bold is not None:
                run.bold = last_run.bold
            if last_run.italic is not None:
                run.italic = last_run.italic

    return True
