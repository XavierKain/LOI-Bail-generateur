"""
Module pour charger les textes formatés depuis un document Word.
Alternative à l'extraction depuis Excel avec rich text.
"""

from pathlib import Path
from typing import Dict, Optional
from docx import Document
import re


class WordTextLoader:
    """Charge les textes formatés depuis un document Word."""

    def __init__(self, word_file_path: str = "Textes BAIL avec styles.docx"):
        """
        Initialise le loader.

        Args:
            word_file_path: Chemin vers le document Word contenant les textes formatés
        """
        self.word_file_path = Path(word_file_path)
        self.sections = {}  # {section_id: paragraph_object}

        if self.word_file_path.exists():
            self._load_sections()

    def _load_sections(self):
        """Charge toutes les sections depuis le document Word."""
        try:
            doc = Document(self.word_file_path)

            current_section_id = None
            skip_next = False

            for para in doc.paragraphs:
                text = para.text.strip()

                # Détecter les identifiants de section [ID: ...]
                if text.startswith('[ID:') and text.endswith(']'):
                    # Extraire l'ID
                    section_id = text[4:-1].strip()
                    current_section_id = section_id
                    skip_next = False
                    continue

                # Skip les lignes de métadonnées (Désignation, Condition)
                if text.startswith('Désignation:') or text.startswith('Condition:'):
                    continue

                # Skip les séparateurs
                if text.startswith('─'):
                    current_section_id = None
                    continue

                # Skip les titres d'article et instructions
                if text.startswith('ARTICLE:') or text.startswith('INSTRUCTIONS:') or text.startswith('TEXTES BAIL'):
                    continue

                # Si on a un ID et du texte, c'est le contenu formaté
                if current_section_id and text and not text.startswith('[ID:'):
                    # Stocker le paragraphe entier (avec son formatage)
                    self.sections[current_section_id] = para
                    # Un seul paragraphe par section
                    current_section_id = None

            print(f"✅ WordTextLoader: {len(self.sections)} sections chargées depuis {self.word_file_path}")

        except Exception as e:
            print(f"⚠️ WordTextLoader: Impossible de charger {self.word_file_path}: {e}")
            self.sections = {}

    def get_formatted_paragraph(self, section_id: str):
        """
        Récupère le paragraphe formaté pour un ID de section.

        Args:
            section_id: L'identifiant de la section (ex: "COMPARUTION_COMPARUTION_BAILLEUR_1")

        Returns:
            Le paragraphe python-docx avec son formatage, ou None si non trouvé
        """
        return self.sections.get(section_id)

    def get_section_ids(self) -> list:
        """Retourne la liste de tous les IDs de section disponibles."""
        return list(self.sections.keys())

    def has_section(self, section_id: str) -> bool:
        """Vérifie si une section existe."""
        return section_id in self.sections

    def copy_formatted_text_to_paragraph(self, source_para, target_para):
        """
        Copie le texte formaté d'un paragraphe source vers un paragraphe cible.

        Args:
            source_para: Paragraphe source (depuis le document Word de styles)
            target_para: Paragraphe cible (dans le document généré)
        """
        # Vider le paragraphe cible
        for run in target_para.runs:
            run.text = ""

        # Copier chaque run avec son formatage
        for source_run in source_para.runs:
            target_run = target_para.add_run(source_run.text)

            # Copier les propriétés de formatage
            if source_run.bold is not None:
                target_run.bold = source_run.bold
            if source_run.italic is not None:
                target_run.italic = source_run.italic
            if source_run.underline is not None:
                target_run.underline = source_run.underline

            # Copier la police
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            if source_run.font.color and source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb

    def __str__(self):
        """Représentation string pour debug."""
        return f"WordTextLoader({len(self.sections)} sections depuis {self.word_file_path})"
