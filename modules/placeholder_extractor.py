"""
Module pour extraire tous les placeholders d'un document Word.
"""

from docx import Document
from pathlib import Path
import re
from typing import Set, List
import logging

logger = logging.getLogger(__name__)


def extract_all_placeholders(template_path: str) -> Set[str]:
    """
    Extrait tous les placeholders [Variable] et {{ARTICLE}} d'un template Word.

    Args:
        template_path: Chemin vers le template Word

    Returns:
        Set de noms de placeholders (sans les crochets)
    """
    placeholders = set()

    try:
        doc = Document(template_path)

        # Extraire depuis les paragraphes
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Trouver [Variable]
            placeholders.update(re.findall(r'\[([^\]]+)\]', text))
            # Trouver {{ARTICLE}}
            placeholders.update(re.findall(r'\{\{([^}]+)\}\}', text))

        # Extraire depuis les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        placeholders.update(re.findall(r'\[([^\]]+)\]', text))
                        placeholders.update(re.findall(r'\{\{([^}]+)\}\}', text))

        logger.info(f"{len(placeholders)} placeholders extraits du template")
        return placeholders

    except Exception as e:
        logger.error(f"Erreur lors de l'extraction des placeholders: {e}")
        return set()


def categorize_placeholders(placeholders: Set[str]) -> dict:
    """
    Catégorise les placeholders par type.

    Args:
        placeholders: Set de placeholders

    Returns:
        Dict avec catégories: articles, variables normales, variables en lettres
    """
    articles = set()
    variables_lettres = set()
    variables_normales = set()

    for p in placeholders:
        if p.startswith("ARTICLE") or p.startswith("COMPARUTION") or p in ["VILLE", "DATE_SIGNATURE"]:
            articles.add(p)
        elif p.endswith(" en lettres"):
            variables_lettres.add(p)
        else:
            variables_normales.add(p)

    return {
        "articles": sorted(articles),
        "variables_lettres": sorted(variables_lettres),
        "variables_normales": sorted(variables_normales)
    }
