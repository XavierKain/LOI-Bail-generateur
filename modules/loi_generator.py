"""
Module pour générer les documents LOI depuis le template Word.
Gère le remplacement des placeholders, les sections optionnelles, et les headers/footers.
"""

import logging
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime, timedelta
from docx import Document
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

logger = logging.getLogger(__name__)


class LOIGenerator:
    """Générateur de documents LOI depuis un template DOCX."""

    def __init__(
        self,
        variables: Dict[str, str],
        societes_info: Dict[str, Dict[str, str]],
        template_path: str = "Template LOI avec placeholder.docx"
    ):
        """
        Initialise le générateur.

        Args:
            variables: Dictionnaire des variables extraites
            societes_info: Informations des sociétés bailleures
            template_path: Chemin vers le template Word
        """
        self.variables = variables.copy()
        self.societes_info = societes_info
        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Template introuvable: {template_path}")

        # Normaliser les noms de variables (mapping)
        self._normalize_variable_names()

        # Pré-calculer les valeurs dérivées
        self._calculate_derived_values()

        logger.info("Générateur LOI initialisé")

    def _normalize_variable_names(self):
        """
        Normalise les noms de variables pour gérer les variations.
        Crée des alias pour les variables avec des noms légèrement différents.
        """
        # Mapping: Nom dans Excel → Nom dans Template
        mappings = {
            "Statut Locaux loués": "Statut Locaux Loués",
            "Durée ferme bail": "Durée ferme Bail",
            "Duré GAPD": "Durée GAPD",  # Typo dans l'ancien Excel
        }

        for excel_name, template_name in mappings.items():
            if excel_name in self.variables and template_name not in self.variables:
                self.variables[template_name] = self.variables[excel_name]
                logger.debug(f"Mapping: '{excel_name}' → '{template_name}'")

        # Créer un mapping case-insensitive pour toutes les variables
        # pour éviter les problèmes de casse
        var_keys = list(self.variables.keys())
        for key in var_keys:
            key_lower = key.lower()
            # Chercher si une version avec casse différente existe
            for other_key in var_keys:
                if other_key != key and other_key.lower() == key_lower:
                    # Utiliser la version avec casse "correcte" (celle du template)
                    if not self.variables.get(other_key):
                        self.variables[other_key] = self.variables[key]
                        logger.debug(f"Case-insensitive mapping: '{key}' → '{other_key}'")

    def _copy_run_format(self, source_run, target_run, override_color=None):
        """
        Copie TOUS les attributs de formatage d'un run source vers un run cible.

        Args:
            source_run: Run source dont on copie le formatage
            target_run: Run cible qui reçoit le formatage
            override_color: Couleur RGB à forcer (pour mettre en rouge/noir)
        """
        if not source_run:
            return

        # Copier tous les attributs de police
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.strike = source_run.font.strike
        target_run.font.subscript = source_run.font.subscript
        target_run.font.superscript = source_run.font.superscript
        target_run.font.all_caps = source_run.font.all_caps
        target_run.font.small_caps = source_run.font.small_caps

        # Couleur: utiliser override_color si fourni, sinon copier
        if override_color:
            target_run.font.color.rgb = override_color
        elif source_run.font.color and source_run.font.color.type == 1:
            target_run.font.color.rgb = source_run.font.color.rgb

    def _calculate_derived_values(self):
        """Calcule les valeurs dérivées (paliers, adresse, type bail, etc.)."""

        # 1. Calculer les paliers (remises)
        self._calculate_paliers()

        # 2. Construire l'adresse complète
        self._build_address()

        # 3. Calculer le type de bail
        self._calculate_type_bail()

        # 4. Calculer la date de signature (Date d'aujourd'hui + 15 jours)
        self._calculate_date_signature()

        # 5. Calculer les surfaces
        self._calculate_surfaces()

    def _calculate_paliers(self):
        """Calcule les montants des paliers (remises) pour chaque année."""
        try:
            loyer_base_str = self.variables.get("Montant du loyer", "0").strip()
            # Nettoyer la valeur (retirer espaces, virgules)
            loyer_base_str = loyer_base_str.replace(" ", "").replace(",", "")

            try:
                loyer_base = float(loyer_base_str)
            except ValueError:
                logger.warning(f"Montant du loyer invalide: {loyer_base_str}")
                loyer_base = 0

            for annee in range(1, 7):
                loyer_annee_key = f"Loyer année {annee}"
                loyer_annee_str = self.variables.get(loyer_annee_key, "").strip()

                if loyer_annee_str:
                    # Nettoyer la valeur
                    loyer_annee_str = loyer_annee_str.replace(" ", "").replace(",", "")
                    try:
                        loyer_annee = float(loyer_annee_str)
                        remise = loyer_base - loyer_annee

                        if remise > 0:
                            # Formater avec espaces pour les milliers
                            self.variables[f"Montant du palier {annee}"] = f"{int(remise):,}".replace(",", " ")
                            logger.debug(f"Palier année {annee}: {int(remise):,} €")
                    except ValueError:
                        logger.warning(f"Loyer année {annee} invalide: {loyer_annee_str}")

        except Exception as e:
            logger.error(f"Erreur calcul des paliers: {e}")

    def _build_address(self):
        """Construit l'adresse complète des locaux loués."""
        ville = self.variables.get("Ville ou arrondissement", "")
        rue = self.variables.get("Numéro et rue", "")

        if ville and rue:
            self.variables["Adresse Locaux Loués"] = f"{rue}, {ville}"
        elif ville:
            self.variables["Adresse Locaux Loués"] = ville
        elif rue:
            self.variables["Adresse Locaux Loués"] = rue

    def _calculate_type_bail(self):
        """Calcule le type de bail selon la durée."""
        duree_bail = self.variables.get("Durée Bail", "").strip()

        if duree_bail:
            try:
                duree = int(float(duree_bail))
                if duree == 9:
                    self.variables["Type Bail"] = "3/6/9"
                elif duree == 10:
                    self.variables["Type Bail"] = "6/9/10"
                else:
                    self.variables["Type Bail"] = f"{duree} ans"
            except ValueError:
                pass

    def _calculate_date_signature(self):
        """Calcule la date de signature (Date d'aujourd'hui + 15 jours)."""
        date_aujourdhui_str = self.variables.get("Date d'aujourd'hui", "")

        if date_aujourdhui_str:
            try:
                # Parser DD/MM/YYYY
                date_aujourdhui = datetime.strptime(date_aujourdhui_str, "%d/%m/%Y")
                date_signature = date_aujourdhui + timedelta(days=15)
                self.variables["Date de signature"] = date_signature.strftime("%d/%m/%Y")
            except ValueError:
                pass

    def _calculate_surfaces(self):
        """Calcule la surface R-1 (total - RDC)."""
        try:
            surface_totale = self.variables.get("Surface totale", "").strip()
            surface_rdc = self.variables.get("Surface RDC", "").strip()

            if surface_totale and surface_rdc:
                surface_totale = float(surface_totale.replace(" ", "").replace(",", "."))
                surface_rdc = float(surface_rdc.replace(" ", "").replace(",", "."))
                surface_r1 = surface_totale - surface_rdc

                if surface_r1 > 0:
                    self.variables["Surface R-1"] = str(int(surface_r1))
        except ValueError:
            pass

    def _is_paragraph_optional(self, paragraph) -> bool:
        """
        Détecte si un paragraphe est optionnel (texte en bleu).

        Args:
            paragraph: Paragraphe docx

        Returns:
            True si le paragraphe est optionnel (bleu)
        """
        for run in paragraph.runs:
            if run.font.color and run.font.color.type == 1:  # RGB color
                rgb = run.font.color.rgb
                # Bleu: B > R et B > G
                # RGBColor est un tuple (R, G, B) indexable
                if rgb and len(rgb) >= 3:
                    r, g, b = rgb[0], rgb[1], rgb[2]
                    if b > r and b > g:
                        return True
        return False

    def _find_placeholders(self, text: str) -> List[str]:
        """
        Trouve tous les placeholders dans un texte.

        Args:
            text: Texte à analyser

        Returns:
            Liste des placeholders trouvés
        """
        return re.findall(r'\[([^\]]+)\]', text)

    def _has_data_for_placeholders(self, placeholders: List[str]) -> bool:
        """
        Vérifie si toutes les données sont disponibles pour les placeholders.

        Args:
            placeholders: Liste des placeholders

        Returns:
            True si toutes les données sont disponibles
        """
        for placeholder in placeholders:
            if placeholder not in self.variables or not self.variables[placeholder]:
                return False
        return True

    def _replace_placeholders_in_text(self, text: str) -> Tuple[str, bool]:
        """
        Remplace les placeholders dans un texte.

        Args:
            text: Texte contenant des placeholders

        Returns:
            Tuple (texte_modifié, données_manquantes)
        """
        placeholders = self._find_placeholders(text)
        missing_data = False

        for placeholder in placeholders:
            value = self.variables.get(placeholder, "")
            if value:
                text = text.replace(f"[{placeholder}]", value)
            else:
                missing_data = True

        return text, missing_data

    def _process_paragraph(self, paragraph) -> Optional[str]:
        """
        Traite un paragraphe: remplace les placeholders ou le supprime.
        Gère les placeholders fragmentés entre plusieurs runs.

        Args:
            paragraph: Paragraphe docx

        Returns:
            "delete" si le paragraphe doit être supprimé, None sinon
        """
        # Utiliser le texte complet du paragraphe (reconstitué depuis tous les runs)
        full_text = paragraph.text
        placeholders = self._find_placeholders(full_text)

        if not placeholders:
            return None

        is_optional = self._is_paragraph_optional(paragraph)
        has_data = self._has_data_for_placeholders(placeholders)

        # Section optionnelle sans données → Supprimer
        if is_optional and not has_data:
            logger.debug(f"Suppression paragraphe optionnel: {full_text[:50]}...")
            return "delete"

        # Section optionnelle avec données → Mettre TOUT en noir
        if is_optional and has_data:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Vérifier si on a des données manquantes
        has_missing_data = False
        for placeholder in placeholders:
            if not self.variables.get(placeholder, ""):
                has_missing_data = True
                break

        import re

        # CAS 1: Section OBLIGATOIRE avec données manquantes → Reconstruire pour mettre les placeholders en rouge
        if not is_optional and has_missing_data:
            # Créer un mapping détaillé: position → (run, formatage)
            char_to_run_map = []
            for run in paragraph.runs:
                for _ in range(len(run.text)):
                    char_to_run_map.append(run)

            # Sauvegarder runs originaux
            original_runs = list(paragraph.runs)

            # Effacer tous les runs
            for run in list(paragraph.runs):
                run.text = ""

            # Construire une liste de segments avec leur formatage
            # Segment = (start_pos, end_pos, text, run_source, is_placeholder, is_missing)
            segments = []
            current_pos = 0

            # Identifier tous les placeholders et leurs positions
            placeholder_matches = list(re.finditer(r'\[([^\]]+)\]', full_text))

            # Si pas de placeholders, juste copier le texte
            if not placeholder_matches:
                pos = 0
                for run in original_runs:
                    if run.text:
                        segments.append((pos, pos + len(run.text), run.text, run, False, False))
                        pos += len(run.text)
            else:
                # Traiter le texte segment par segment en respectant les changements de run ET les placeholders
                pos = 0
                placeholder_idx = 0

                while pos < len(full_text):
                    # Y a-t-il un placeholder à cette position ?
                    if placeholder_idx < len(placeholder_matches):
                        match = placeholder_matches[placeholder_idx]
                        ph_start, ph_end = match.span()
                        placeholder = match.group(1)

                        # Si on est avant le placeholder, ajouter le texte normal
                        if pos < ph_start:
                            # Découper par changement de run
                            current_run = char_to_run_map[pos] if pos < len(char_to_run_map) else original_runs[0]
                            segment_start = pos
                            while pos < ph_start and pos < len(char_to_run_map):
                                if char_to_run_map[pos] != current_run:
                                    # Changement de run: sauvegarder le segment
                                    segments.append((segment_start, pos, full_text[segment_start:pos], current_run, False, False))
                                    segment_start = pos
                                    current_run = char_to_run_map[pos]
                                pos += 1
                            # Sauvegarder le dernier segment avant le placeholder
                            if segment_start < pos:
                                segments.append((segment_start, pos, full_text[segment_start:pos], current_run, False, False))

                        # Ajouter le placeholder ou sa valeur
                        value = self.variables.get(placeholder, "")
                        source_run = char_to_run_map[ph_start] if ph_start < len(char_to_run_map) else original_runs[0]
                        if value:
                            segments.append((ph_start, ph_end, value, source_run, True, False))
                        else:
                            segments.append((ph_start, ph_end, f"[{placeholder}]", source_run, True, True))

                        pos = ph_end
                        placeholder_idx += 1
                    else:
                        # Plus de placeholders, traiter le reste
                        if pos < len(full_text):
                            current_run = char_to_run_map[pos] if pos < len(char_to_run_map) else original_runs[0]
                            segment_start = pos
                            while pos < len(char_to_run_map):
                                if char_to_run_map[pos] != current_run:
                                    segments.append((segment_start, pos, full_text[segment_start:pos], current_run, False, False))
                                    segment_start = pos
                                    current_run = char_to_run_map[pos]
                                pos += 1
                            if segment_start < pos:
                                segments.append((segment_start, pos, full_text[segment_start:pos], current_run, False, False))
                        break

            # Reconstruire en préservant le formatage de chaque segment
            for _, _, text, source_run, is_ph, is_missing in segments:
                if text:
                    new_run = paragraph.add_run(text)
                    if is_missing:
                        override_color = RGBColor(255, 0, 0)
                    else:
                        override_color = RGBColor(0, 0, 0)
                    self._copy_run_format(source_run, new_run, override_color=override_color)

            if has_missing_data:
                logger.warning(f"Placeholder manquant (rouge): {full_text[:50]}...")

        # CAS 2: Toutes les données présentes → Remplacer dans le texte SANS toucher au formatage
        else:
            # Remplacer les placeholders dans chaque run INDIVIDUELLEMENT
            for placeholder in placeholders:
                value = self.variables.get(placeholder, "")
                if value:
                    placeholder_pattern = f"[{placeholder}]"
                    for run in paragraph.runs:
                        if placeholder_pattern in run.text:
                            run.text = run.text.replace(placeholder_pattern, value)
                            # Le formatage du run est automatiquement préservé

        return None

    def _update_headers_footers(self, doc: Document):
        """
        Met à jour les headers et footers selon la société bailleur.
        Ne touche PAS à la structure existante, remplace juste le contenu texte.

        Args:
            doc: Document docx
        """
        societe_bailleur = self.variables.get("Société Bailleur", "")

        if not societe_bailleur or societe_bailleur not in self.societes_info:
            logger.warning(f"Société bailleur '{societe_bailleur}' non trouvée dans la config")
            return

        societe_info = self.societes_info[societe_bailleur]
        header_text = societe_info.get("header", "")
        footer_text = societe_info.get("footer", "")

        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Mettre à jour tous les headers et footers
        for section in doc.sections:
            # Header
            if header_text:
                header = section.header

                # Ajuster les marges du header
                # Réduire la marge supérieure et augmenter la marge inférieure
                section.top_margin = Inches(0.5)  # Marge haut de page réduite
                section.header_distance = Inches(0.3)  # Distance du header au texte

                # Supprimer tous les paragraphes existants
                for para in list(header.paragraphs):
                    p = para._element
                    p.getparent().remove(p)

                # Créer un nouveau paragraphe avec le bon formatage
                para = header.add_paragraph()
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.add_run(header_text)
                run.font.bold = True
                run.font.size = Pt(22)  # Police 22pt au lieu de 12pt

                # Ajouter un espacement après le header
                from docx.shared import Pt as PtSpace
                para.paragraph_format.space_after = PtSpace(12)

            # Footer
            if footer_text:
                footer = section.footer

                # Ajuster les marges du footer
                # Augmenter la marge supérieure (espace avant le footer) et réduire la marge inférieure
                section.bottom_margin = Inches(0.5)  # Marge bas de page réduite
                section.footer_distance = Inches(0.3)  # Distance du texte au footer

                # Supprimer tous les paragraphes existants
                for para in list(footer.paragraphs):
                    p = para._element
                    p.getparent().remove(p)

                # Le footer peut avoir plusieurs lignes
                lines = footer_text.split("\n")
                for i, line in enumerate(lines):
                    para = footer.add_paragraph()
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = para.add_run(line)
                    run.font.size = Pt(9)

                    # Ajouter un espacement avant le premier paragraphe du footer
                    if i == 0:
                        from docx.shared import Pt as PtSpace
                        para.paragraph_format.space_before = PtSpace(12)

        logger.info(f"Headers/Footers mis à jour pour: {societe_bailleur}")

    def generate(self, output_path: str) -> str:
        """
        Génère le document LOI final.

        Args:
            output_path: Chemin du fichier de sortie

        Returns:
            Chemin du fichier généré
        """
        logger.info(f"Génération du document LOI: {output_path}")

        # Charger le template
        doc = Document(str(self.template_path))

        # Première passe: identifier les sections à garder
        # Pour les paragraphes comme "Remises..." qui n'ont pas de placeholder mais contrôlent une section
        paragraphs_with_data = set()
        all_paragraphs = doc.paragraphs

        # Détecter si les paliers ont des données
        has_palier_data = any(
            self.variables.get(f"Montant du palier {i}", "")
            for i in range(1, 7)
        )

        # Détecter si les conditions suspensives ont des données
        has_conditions_data = any(
            self.variables.get(f"Condition suspensive {i}", "")
            for i in range(1, 5)
        )

        # Traiter tous les paragraphes
        paragraphs_to_delete = []
        for i, paragraph in enumerate(all_paragraphs):
            text = paragraph.text
            is_optional = self._is_paragraph_optional(paragraph)
            placeholders = self._find_placeholders(text)

            # Cas spéciaux: paragraphes de titre sans placeholder OU avec placeholder [.]
            if is_optional:
                # "Remises (sur loyer annuel indexé) :" - sans placeholder
                if "Remises" in text and "loyer" in text and not placeholders:
                    if has_palier_data:
                        # Garder ce paragraphe et le mettre en noir
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        paragraphs_to_delete.append(paragraph)
                    continue

                # "Condition(s) suspensive(s) : à réaliser au plus tard pour le [.]"
                if "Condition" in text and "suspensive" in text and "[.]" in text:
                    if has_conditions_data:
                        # Garder ce paragraphe et le mettre en noir (mais laisser le traitement normal faire le remplacement)
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        # Ne pas faire continue - laisser le traitement normal gérer les placeholders
                    else:
                        paragraphs_to_delete.append(paragraph)
                        continue

                # "Franchise de loyer : [Durée Franchise]..." et "Garantie à première demande... [Durée GAPD]..."
                # Ces paragraphes ont des placeholders et doivent être traités normalement
                # mais on s'assure que TOUS les runs passent en noir si on a les données
                if placeholders:
                    has_data = self._has_data_for_placeholders(placeholders)
                    if has_data:
                        # Mettre TOUS les runs en noir
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)

            # Traitement normal
            result = self._process_paragraph(paragraph)
            if result == "delete":
                paragraphs_to_delete.append(paragraph)

        # Supprimer les paragraphes marqués
        for paragraph in paragraphs_to_delete:
            p = paragraph._element
            p.getparent().remove(p)

        # Traiter les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cells_to_delete = []
                    for paragraph in cell.paragraphs:
                        result = self._process_paragraph(paragraph)
                        if result == "delete":
                            cells_to_delete.append(paragraph)

                    # Supprimer les paragraphes dans les cellules
                    for paragraph in cells_to_delete:
                        p = paragraph._element
                        p.getparent().remove(p)

        # Mettre à jour les headers/footers
        self._update_headers_footers(doc)

        # Sauvegarder
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        logger.info(f"Document généré: {output_path}")
        return str(output_path)
