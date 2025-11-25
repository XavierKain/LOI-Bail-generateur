"""
Module de génération de documents BAIL au format Word.

Ce module prend les articles générés par BailGenerator et les insère
dans le template Word avec placeholders, puis remplace tous les placeholders
[Variable] par les valeurs extraites.
"""

from docx import Document
from docx.shared import RGBColor, Pt
from pathlib import Path
from typing import Dict
import logging
import re
from .number_to_french import number_to_french_words

logger = logging.getLogger(__name__)

# Police par défaut pour le contenu du BAIL
DEFAULT_FONT_NAME = "Calibri"
DEFAULT_FONT_SIZE = Pt(11)


class BailWordGenerator:
    """Générateur de documents BAIL au format Word."""

    @staticmethod
    def _apply_default_font(run):
        """Applique la police par défaut (Calibri 11) à un run."""
        run.font.name = DEFAULT_FONT_NAME
        run.font.size = DEFAULT_FONT_SIZE

    @staticmethod
    def _parse_formatting_tags(text: str) -> list:
        """
        Parse le texte contenant des balises HTML-like (<b>, <i>, <u>) et retourne
        une liste de segments avec leur formatage.

        Args:
            text: Texte contenant potentiellement des balises <b>, <i>, <u>

        Returns:
            Liste de tuples (texte, {formatage}) où formatage = {"bold": bool, "italic": bool, "underline": bool}

        Example:
            "La <b>Société [Nom]</b> est <i>présente</i>"
            -> [("La ", {}), ("Société [Nom]", {"bold": True}), (" est ", {}), ("présente", {"italic": True})]
        """
        if not text or not any(tag in text for tag in ['<b>', '<i>', '<u>']):
            # Pas de balises, retourner le texte tel quel
            return [(text, {})]

        segments = []
        current_pos = 0
        format_stack = []  # Stack pour gérer les balises imbriquées

        # Pattern pour trouver toutes les balises ouvrantes et fermantes
        tag_pattern = re.compile(r'<(/?)([biu])>', re.IGNORECASE)

        # Trouver toutes les balises
        matches = list(tag_pattern.finditer(text))

        if not matches:
            return [(text, {})]

        for match in matches:
            # Texte avant la balise
            if match.start() > current_pos:
                text_before = text[current_pos:match.start()]
                # Appliquer le formatage actuel du stack
                current_format = {}
                for fmt in format_stack:
                    current_format[fmt] = True
                segments.append((text_before, current_format.copy()))

            # Traiter la balise
            is_closing = match.group(1) == '/'
            tag_type = match.group(2).lower()

            # Mapper le tag au nom de propriété
            tag_map = {'b': 'bold', 'i': 'italic', 'u': 'underline'}
            format_name = tag_map.get(tag_type)

            if is_closing:
                # Balise fermante: retirer du stack
                if format_name in format_stack:
                    format_stack.remove(format_name)
            else:
                # Balise ouvrante: ajouter au stack
                if format_name not in format_stack:
                    format_stack.append(format_name)

            current_pos = match.end()

        # Texte après la dernière balise
        if current_pos < len(text):
            text_after = text[current_pos:]
            current_format = {}
            for fmt in format_stack:
                current_format[fmt] = True
            segments.append((text_after, current_format.copy()))

        return segments

    @staticmethod
    def _apply_formatting(run, formatting: dict):
        """
        Applique le formatage (bold, italic, underline) à un run.

        Args:
            run: Run docx
            formatting: Dict avec keys "bold", "italic", "underline" (valeurs bool)
        """
        if formatting.get('bold'):
            run.font.bold = True
        if formatting.get('italic'):
            run.font.italic = True
        if formatting.get('underline'):
            run.font.underline = True

    def _normalize_variable_name(self, var_name: str, donnees: Dict[str, any]) -> str:
        """
        Normalise le nom de variable pour gérer les variations.

        Args:
            var_name: Nom de variable brut
            donnees: Données disponibles

        Returns:
            Nom de variable normalisé ou original si trouvé directement
        """
        # Si la variable existe telle quelle, la retourner
        if var_name in donnees:
            return var_name

        # Mappings courants
        mappings = {
            # Montant Palier X → Montant du palier X
            **{f"Montant Palier {i}": f"Montant du palier {i}" for i in range(1, 7)},
            **{f"Montant palier {i}": f"Montant du palier {i}" for i in range(1, 7)},
        }

        normalized = mappings.get(var_name, var_name)

        # Si toujours pas trouvé, chercher avec des variations de casse
        if normalized not in donnees:
            for key in donnees.keys():
                if key.lower() == normalized.lower():
                    return key

        return normalized

    def __init__(self, template_path: str = "2025 - Template BAIL.docx"):
        """
        Initialise le générateur Word pour BAIL.

        Args:
            template_path: Chemin vers le template Word avec placeholders
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")

        logger.info(f"Template BAIL chargé: {template_path}")

    def generer_document(
        self,
        articles_generes: Dict[str, str],
        donnees: Dict[str, any],
        output_path: str
    ) -> None:
        """
        Génère le document BAIL Word final.

        Args:
            articles_generes: Dict avec les articles générés par BailGenerator
            donnees: Données complètes (variables extraites + dérivées)
            output_path: Chemin de sortie pour le document généré

        Raises:
            FileNotFoundError: Si le template n'existe pas
        """
        logger.info("Début de la génération du document BAIL Word")

        # Charger le template
        doc = Document(self.template_path)

        # ÉTAPE 1: Remplacer les placeholders {{ARTICLE}}
        placeholder_mapping = {
            "{{COMPARUTION_BAILLEUR}}": articles_generes.get("Comparution Bailleur", ""),
            "{{COMPARUTION_PRENEUR}}": articles_generes.get("Comparution Preneur", ""),
            "{{ARTICLE_PRELIMINAIRE}}": articles_generes.get("Article préliminaire", ""),
            "{{ARTICLE_1}}": articles_generes.get("Article 1", ""),
            "{{ARTICLE_2}}": articles_generes.get("Article 2", ""),
            "{{ARTICLE_3}}": articles_generes.get("Article 3", ""),
            "{{ARTICLE_5_3}}": articles_generes.get("Article 5.3", ""),
            "{{ARTICLE_7_1}}": articles_generes.get("Article 7.1", ""),
            "{{ARTICLE_7_2}}": articles_generes.get("Article 7.2", ""),
            "{{ARTICLE_7_3}}": articles_generes.get("Article  7.3", ""),  # Note: 2 espaces
            "{{ARTICLE_7_6}}": articles_generes.get("Article 7.6", ""),
            "{{ARTICLE_8}}": articles_generes.get("Article 8", ""),
            "{{ARTICLE_19}}": articles_generes.get("Article 19", ""),
            "{{ARTICLE_22_2}}": articles_generes.get("Article 22.2", ""),
            "{{ARTICLE_26}}": articles_generes.get("Article 26", ""),
            "{{ARTICLE_26_1}}": articles_generes.get("Article 26.1", ""),
            "{{ARTICLE_26_2}}": articles_generes.get("Article 26.2", ""),
            "{{VILLE}}": donnees.get("Ville ou arrondissement", "Paris").split("(")[0].strip(),
            "{{DATE_SIGNATURE}}": donnees.get("Date de signature", ""),
        }

        # Remplacer les placeholders {{ARTICLE}} dans tous les paragraphes
        for paragraph in doc.paragraphs:
            self._replace_article_placeholders(paragraph, placeholder_mapping, doc)

        # Remplacer les placeholders dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_article_placeholders(paragraph, placeholder_mapping, doc)

        # ÉTAPE 2: Remplacer les placeholders [Variable] dans TOUT le document
        # (comme dans LOIGenerator)
        for paragraph in doc.paragraphs:
            self._replace_variable_placeholders(paragraph, donnees)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_variable_placeholders(paragraph, donnees)

        # Nettoyer uniquement les placeholders {{}} non remplacés
        self._clean_unreplaced_placeholders(doc)

        # Corriger l'indentation des headings
        self._fix_heading_indentation(doc)

        # Mettre à jour la table des matières
        self._update_toc(doc)

        # Sauvegarder le document
        doc.save(output_path)
        logger.info(f"Document BAIL généré: {output_path}")

    def _get_comparution_bailleur(self, articles: Dict[str, str]) -> str:
        """
        Extrait le texte de comparution du Bailleur.

        Args:
            articles: Articles générés

        Returns:
            Texte de comparution du Bailleur
        """
        comparution = articles.get("Comparution", "")
        if not comparution:
            return ""

        # Le texte complet contient bailleur et preneur
        # On cherche la partie bailleur (avant "D'UNE PART")
        parts = comparution.split("D'UNE PART")
        if len(parts) >= 1:
            return parts[0].strip()
        return comparution

    def _get_comparution_preneur(self, articles: Dict[str, str]) -> str:
        """
        Extrait le texte de comparution du Preneur.

        Args:
            articles: Articles générés

        Returns:
            Texte de comparution du Preneur
        """
        comparution = articles.get("Comparution", "")
        if not comparution:
            return ""

        # Le texte complet contient bailleur et preneur
        # On cherche la partie preneur (après "ET :" et avant "D'AUTRE PART")
        if "ET :" in comparution:
            parts = comparution.split("ET :")
            if len(parts) >= 2:
                preneur_part = parts[1].split("D'AUTRE PART")[0] if "D'AUTRE PART" in parts[1] else parts[1]
                return preneur_part.strip()

        return ""

    def _replace_article_placeholders(
        self,
        paragraph,
        mapping: Dict[str, str],
        doc=None
    ) -> None:
        """
        Remplace les placeholders {{ARTICLE}} dans un paragraphe.
        Parse et applique les balises de formatage HTML-like (<b>, <i>, <u>).
        Gère les marqueurs de titre (** pour Heading 2, *** pour Heading 3).
        Si le texte contient plusieurs lignes avec marqueurs, crée des paragraphes séparés.

        Args:
            paragraph: Paragraphe docx
            mapping: Mapping {placeholder: texte_final}
            doc: Document docx (optionnel, nécessaire pour créer de nouveaux paragraphes)
        """
        full_text = paragraph.text

        # Vérifier s'il y a des placeholders {{}}
        if "{{" not in full_text:
            return

        # Pour chaque placeholder trouvé
        for placeholder, replacement in mapping.items():
            if placeholder in full_text:
                # Si le remplacement est vide, on supprime le placeholder
                if not replacement:
                    full_text = full_text.replace(placeholder, "")
                else:
                    full_text = full_text.replace(placeholder, replacement)

        # Si le texte a changé, on met à jour le paragraphe
        if full_text != paragraph.text:
            # Stratégie améliorée : diviser d'abord sur \n\n, puis sur \n si on trouve des marqueurs ** ou ***
            # Ceci gère les cas où l'Excel a des sauts de ligne incohérents

            # Stratégie simple: Split sur \n\n et insérer un paragraphe vide entre chaque partie
            import re

            # Split sur \n\n
            parts_raw = full_text.split('\n\n')

            # Créer final_paragraphs en intercalant des éléments vides
            final_paragraphs = []
            for i, part in enumerate(parts_raw):
                # Traiter les headings dans cette partie
                heading_parts = re.split(r'\n(?=\*{2,4})', part)

                for hp in heading_parts:
                    stripped = hp.strip()
                    if stripped:  # Ne pas ajouter si complètement vide
                        final_paragraphs.append(stripped)

                # Ajouter un paragraphe vide APRÈS (sauf pour le dernier)
                if i < len(parts_raw) - 1:
                    final_paragraphs.append('')  # Paragraphe vide pour l'espacement


            if not final_paragraphs:
                return

            # Traiter le premier paragraphe dans le paragraphe Word actuel
            self._process_paragraph_with_heading(paragraph, final_paragraphs[0])

            # Pour les paragraphes suivants, créer de nouveaux paragraphes Word si doc est fourni
            if doc and len(final_paragraphs) > 1:
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph

                # Insérer les nouveaux paragraphes après le paragraphe actuel
                last_para_element = paragraph._element

                for para_text in final_paragraphs[1:]:
                    # Créer un nouvel élément de paragraphe directement
                    new_p_element = OxmlElement('w:p')

                    # Insérer après le dernier paragraphe traité
                    last_para_element.addnext(new_p_element)

                    # Créer l'objet Paragraph python-docx
                    new_para = Paragraph(new_p_element, paragraph._parent)

                    # Si le paragraphe est vide, le laisser vide (espacement)
                    if not para_text:
                        # Paragraphe vide - ne rien faire, juste le style Normal
                        try:
                            new_para.style = 'Normal'
                        except:
                            pass
                    else:
                        # Paragraphe avec contenu
                        # Réinitialiser le style à Normal (évite l'héritage du style Heading)
                        try:
                            new_para.style = 'Normal'
                        except:
                            pass

                        # Traiter le paragraphe
                        self._process_paragraph_with_heading(new_para, para_text)

                    # Mettre à jour le dernier élément traité
                    last_para_element = new_p_element

    def _process_paragraph_with_heading(self, paragraph, text: str) -> None:
        """
        Traite un paragraphe unique en détectant les marqueurs de titre et en appliquant le formatage.

        Args:
            paragraph: Paragraphe docx
            text: Texte à traiter
        """
        # Détecter les marqueurs de titre au début du texte
        heading_style = None
        text_to_parse = text

        # Chercher du plus spécifique au moins spécifique (**** avant *** avant **)
        if text.startswith('****'):
            heading_style = 'Heading 4'
            text_to_parse = text[4:].lstrip()  # Retirer **** et espaces
        elif text.startswith('***'):
            heading_style = 'Heading 3'
            text_to_parse = text[3:].lstrip()  # Retirer *** et espaces
        elif text.startswith('**'):
            heading_style = 'Heading 2'
            text_to_parse = text[2:].lstrip()  # Retirer ** et espaces

        # Appliquer le style de titre si détecté
        if heading_style:
            try:
                paragraph.style = heading_style
                # Réinitialiser l'indentation pour aligner à gauche
                paragraph.paragraph_format.left_indent = None
                paragraph.paragraph_format.first_line_indent = None
            except:
                # Si le style n'existe pas, ignorer
                pass

        # Parser les balises de formatage
        segments = self._parse_formatting_tags(text_to_parse)

        # Vider tous les runs existants
        for run in list(paragraph.runs):
            run.text = ""

        # Créer un run pour chaque segment avec son formatage
        for text, formatting in segments:
            if text:  # Ignorer les segments vides
                run = paragraph.add_run(text)
                self._apply_default_font(run)
                self._apply_formatting(run, formatting)

    def _replace_variable_placeholders(
        self,
        paragraph,
        donnees: Dict[str, any]
    ) -> None:
        """
        Remplace les placeholders [Variable] dans un paragraphe.
        Met les placeholders manquants en ROUGE.
        Gère les placeholders "en lettres" pour conversion numérique.
        PRÉSERVE le formatage existant des runs (bold, italic, etc.).

        Args:
            paragraph: Paragraphe docx
            donnees: Données avec toutes les variables
        """
        # Créer le mapping placeholder -> (valeur, is_red)
        placeholder_mapping = {}

        # Trouver tous les placeholders dans le paragraphe complet
        full_text = paragraph.text
        placeholders = re.findall(r'\[([^\]]+)\]', full_text)

        if not placeholders:
            return

        # Pour chaque placeholder, déterminer sa valeur de remplacement et sa couleur
        for placeholder in placeholders:
            # Gestion spéciale pour les placeholders "en lettres"
            if placeholder.endswith(" en lettres"):
                base_variable = placeholder.replace(" en lettres", "")
                base_variable = self._normalize_variable_name(base_variable, donnees)
                value = donnees.get(base_variable)

                if value and str(value).strip():
                    try:
                        value_clean = str(value).replace(" ", "").replace(",", ".")
                        numeric_value = float(value_clean)
                        words = number_to_french_words(numeric_value)
                        placeholder_mapping[f"[{placeholder}]"] = (words + " ", False)
                    except (ValueError, TypeError) as e:
                        logger.warning(f"Impossible de convertir '{value}' en lettres: {e}")
                        placeholder_mapping[f"[{placeholder}]"] = (f"[{placeholder}]", True)
                else:
                    placeholder_mapping[f"[{placeholder}]"] = (f"[{placeholder}]", True)
            else:
                # Placeholder normal
                normalized_placeholder = self._normalize_variable_name(placeholder, donnees)
                value = donnees.get(normalized_placeholder)

                if value and str(value).strip():
                    placeholder_mapping[f"[{placeholder}]"] = (str(value), False)
                else:
                    placeholder_mapping[f"[{placeholder}]"] = (f"[{placeholder}]", True)

        # Maintenant, parcourir chaque run et remplacer les placeholders
        # On crée de nouveaux runs pour chaque remplacement afin de pouvoir colorer individuellement
        runs_to_process = list(paragraph.runs)  # Copie car on va modifier pendant l'itération

        for run in runs_to_process:
            run_text = run.text

            # Vérifier si ce run contient des placeholders
            has_placeholder = any(ph_key in run_text for ph_key in placeholder_mapping.keys())

            if not has_placeholder:
                continue

            # Sauvegarder le formatage du run original
            original_bold = run.font.bold
            original_italic = run.font.italic
            original_underline = run.font.underline
            original_font_name = run.font.name
            original_font_size = run.font.size

            # Diviser le texte du run en segments (texte normal / placeholder)
            segments = []
            remaining = run_text

            while remaining:
                # Trouver le premier placeholder
                first_pos = len(remaining)
                first_placeholder = None

                for ph_key in placeholder_mapping.keys():
                    pos = remaining.find(ph_key)
                    if pos != -1 and pos < first_pos:
                        first_pos = pos
                        first_placeholder = ph_key

                if first_placeholder:
                    # Ajouter le texte avant le placeholder
                    if first_pos > 0:
                        segments.append((remaining[:first_pos], None))  # None = pas un placeholder

                    # Ajouter le placeholder
                    value, is_red = placeholder_mapping[first_placeholder]
                    segments.append((value, is_red))  # is_red = True/False

                    remaining = remaining[first_pos + len(first_placeholder):]
                else:
                    # Plus de placeholders
                    segments.append((remaining, None))
                    break

            # Supprimer le run original
            run._element.getparent().remove(run._element)

            # Créer de nouveaux runs pour chaque segment
            for text, is_red in segments:
                if not text:
                    continue

                new_run = paragraph.add_run(text)

                # Appliquer le formatage du run original
                new_run.font.bold = original_bold
                new_run.font.italic = original_italic
                new_run.font.underline = original_underline
                if original_font_name:
                    new_run.font.name = original_font_name
                if original_font_size:
                    new_run.font.size = original_font_size

                # Appliquer Calibri 11 par défaut
                self._apply_default_font(new_run)

                # Appliquer la couleur (rouge si missing placeholder)
                if is_red:
                    new_run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    new_run.font.color.rgb = RGBColor(0, 0, 0)

    def _update_toc(self, doc) -> None:
        """
        Force la mise à jour de la TOC en recréant le champ TOC.
        Supprime l'ancienne TOC et en crée une nouvelle avec les titres actuels.

        Args:
            doc: Document docx
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement, parse_xml

        # Chercher et supprimer l'ancienne TOC
        toc_found = False
        toc_start_para = None
        toc_end_para = None

        for i, paragraph in enumerate(doc.paragraphs):
            para_xml = paragraph._element.xml.decode('utf-8') if isinstance(paragraph._element.xml, bytes) else str(paragraph._element.xml)

            # Chercher le début de la TOC (TOC \o)
            if 'TOC' in para_xml and not toc_found:
                toc_found = True
                toc_start_para = i
                logger.info(f"TOC trouvée au paragraphe {i}")

            # Chercher la fin de la TOC (fldCharType="end")
            if toc_found and toc_end_para is None:
                if 'fldCharType="end"' in para_xml or 'fldCharType="separate"' in para_xml:
                    # Continuer jusqu'à trouver vraiment la fin
                    continue
                elif i > toc_start_para + 1 and 'fldChar' not in para_xml:
                    # Premier paragraphe après la TOC sans balise de champ
                    toc_end_para = i - 1
                    logger.info(f"Fin de TOC au paragraphe {toc_end_para}")
                    break

        # Marquer tous les champs comme dirty et activer updateFields
        # Cela demandera confirmation à l'ouverture, mais mettra à jour la TOC si l'utilisateur accepte
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for fld_char in run._element.findall(qn('w:fldChar')):
                    fld_type = fld_char.get(qn('w:fldCharType'))
                    if fld_type == 'begin':
                        fld_char.set(qn('w:dirty'), '1')

        # Activer updateFields dans les settings
        try:
            settings_element = doc.settings.element
            update_fields = settings_element.find(qn('w:updateFields'))

            if update_fields is None:
                update_fields = OxmlElement('w:updateFields')
                update_fields.set(qn('w:val'), 'true')
                settings_element.append(update_fields)
            else:
                update_fields.set(qn('w:val'), 'true')

            logger.info("TOC: Le document demandera la mise à jour à l'ouverture")
        except Exception as e:
            logger.warning(f"Impossible de configurer updateFields: {e}")

    def _fix_heading_indentation(self, doc) -> None:
        """
        Réinitialise l'indentation et l'alignement de tous les Headings pour les aligner à gauche.

        Corrige les Heading 4 du template qui ont parfois un first_line_indent ou un alignement centré.
        Nettoie aussi les runs vides qui peuvent causer des décalages visuels.

        Args:
            doc: Document docx
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        fixed_count = 0
        cleaned_runs = 0
        fixed_alignment = 0

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            # Pour tous les Headings (2, 3, 4)
            if style_name.startswith('Heading'):
                # Vérifier s'il y a une indentation
                if (paragraph.paragraph_format.left_indent is not None or
                    paragraph.paragraph_format.first_line_indent is not None):
                    # Réinitialiser
                    paragraph.paragraph_format.left_indent = None
                    paragraph.paragraph_format.first_line_indent = None
                    fixed_count += 1

                # Forcer l'alignement à gauche (désactiver centrage, justifié, etc.)
                if paragraph.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    fixed_alignment += 1

                # Nettoyer les runs vides (peuvent causer des décalages visuels)
                runs_to_remove = []
                for run in paragraph.runs:
                    if len(run.text) == 0:
                        runs_to_remove.append(run)

                for run in runs_to_remove:
                    run._element.getparent().remove(run._element)
                    cleaned_runs += 1

        if fixed_count > 0:
            logger.info(f"Corrigé l'indentation de {fixed_count} headings")
        if fixed_alignment > 0:
            logger.info(f"Corrigé l'alignement de {fixed_alignment} headings (forcé à gauche)")
        if cleaned_runs > 0:
            logger.info(f"Nettoyé {cleaned_runs} runs vides dans les headings")

    def _clean_unreplaced_placeholders(self, doc) -> None:
        """
        Supprime UNIQUEMENT les paragraphes contenant des placeholders {{}} non remplacés.

        NE supprime PAS:
        - Les paragraphes vides (espacement)
        - Les paragraphes Heading vides du template

        Args:
            doc: Document docx
        """
        paragraphs_to_remove = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Supprimer uniquement les paragraphes avec des placeholders non remplacés
            if text and re.match(r'^(\{\{[^}]*\}\}\s*)+$', text):
                paragraphs_to_remove.append(paragraph)
                logger.debug(f"Marqué pour suppression: {text}")

        # Supprimer les paragraphes identifiés
        for paragraph in paragraphs_to_remove:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        if len(paragraphs_to_remove) > 0:
            logger.info(f"Supprimé {len(paragraphs_to_remove)} placeholders non remplacés")

    def _clean_empty_paragraphs(self, doc) -> None:
        """
        DEPRECATED - Utiliser _clean_unreplaced_placeholders à la place.

        Nettoie uniquement les paragraphes qui ne contiennent que des placeholders vides
        ou qui sont vides avec un style Heading.

        NE supprime PAS les paragraphes Normal vides car ils servent d'espacement.

        Args:
            doc: Document docx
        """
        paragraphs_to_remove = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else ""

            # Cas 1: Paragraphe contenant des {{ }} non remplacés
            if text and re.match(r'^(\{\{[^}]*\}\}\s*)+$', text):
                paragraphs_to_remove.append(paragraph)
            # Cas 2: Paragraphe complètement vide avec un style Heading (erreur de formatage)
            elif not text and style_name.startswith('Heading'):
                paragraphs_to_remove.append(paragraph)

        # Supprimer les paragraphes identifiés
        for paragraph in paragraphs_to_remove:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        if len(paragraphs_to_remove) > 0:
            logger.info(f"Nettoyé {len(paragraphs_to_remove)} paragraphes vides")
