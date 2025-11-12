"""Examine signature table structure in detail"""

from docx import Document

template_path = "Template BAIL avec placeholder.docx"
doc = Document(template_path)

# Le tableau de signatures (premier tableau)
table = doc.tables[0]

print("=" * 80)
print("STRUCTURE DU TABLEAU DE SIGNATURES")
print("=" * 80)

print(f"\nNombre de lignes: {len(table.rows)}")
print(f"Nombre de colonnes: {len(table.columns)}")

print("\n📋 Contenu détaillé:")
print("-" * 80)

for row_idx, row in enumerate(table.rows):
    print(f"\nLigne {row_idx}:")
    for cell_idx, cell in enumerate(row.cells):
        print(f"  Colonne {cell_idx}:")
        for para_idx, para in enumerate(cell.paragraphs):
            text = para.text.strip()
            if text:
                print(f"    Paragraphe {para_idx}: '{text}'")
            else:
                print(f"    Paragraphe {para_idx}: (vide)")

print("\n" + "=" * 80)
print("SOLUTION:")
print("=" * 80)
print("""
Le tableau contient 1 ligne avec 2 colonnes:
  - Colonne 0: "Le Bailleur"
  - Colonne 1: "Le Preneur"

Pour ajouter les noms des présidents:
  - Sous "Le Bailleur": ajouter "Monsieur Maxime FORGEOT"
  - Sous "Le Preneur": ajouter le placeholder [PRESIDENT DE LA SOCIETE]

Deux options:
  1. Modifier le template manuellement pour ajouter ces lignes
  2. Modifier le code pour ajouter dynamiquement ces lignes lors de la génération
""")
