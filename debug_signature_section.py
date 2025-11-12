"""Debug signature section in BAIL template"""

from docx import Document

template_path = "Template BAIL avec placeholder.docx"
doc = Document(template_path)

print("=" * 80)
print("RECHERCHE DE LA SECTION SIGNATURES")
print("=" * 80)

# Chercher dans les derniers paragraphes du document
print("\n📄 Derniers 30 paragraphes du document:")
print("-" * 80)

total_paragraphs = len(doc.paragraphs)
start_idx = max(0, total_paragraphs - 30)

for idx in range(start_idx, total_paragraphs):
    para = doc.paragraphs[idx]
    text = para.text.strip()

    if text:  # Seulement les paragraphes non vides
        # Marquer les lignes importantes
        marker = ""
        if any(keyword in text.lower() for keyword in ["bailleur", "preneur", "signature", "fait à", "fait a"]):
            marker = "🎯 "

        print(f"{marker}Paragraphe {idx}: {text[:100]}")

# Chercher aussi dans les tableaux de fin de document
print("\n\n📊 Tableaux dans le document:")
print("-" * 80)

for table_idx, table in enumerate(doc.tables):
    print(f"\nTableau {table_idx}:")

    # Afficher le contenu du tableau
    for row_idx, row in enumerate(table.rows):
        row_text = " | ".join([cell.text.strip() for cell in row.cells])
        if row_text.strip():
            marker = ""
            if any(keyword in row_text.lower() for keyword in ["bailleur", "preneur", "signature"]):
                marker = "🎯 "
            print(f"{marker}  Row {row_idx}: {row_text}")

print("\n" + "=" * 80)
