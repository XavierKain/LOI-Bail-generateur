"""
Test du système de formatage des placeholders.
"""

from docx import Document
from docx.shared import Pt
from modules.placeholder_formatter import replace_placeholder_with_format

print('=' * 80)
print('TEST DU FORMATAGE DES PLACEHOLDERS')
print('=' * 80)

# Créer un document de test
doc = Document()

# Test 1: Placeholder en gras
print('\n1️⃣  Test placeholder en gras')
para1 = doc.add_paragraph()
run1 = para1.add_run("Le preneur est ")
run2 = para1.add_run("[NOM DU PRENEUR]")
run2.bold = True  # Mettre en gras
run3 = para1.add_run(" domicilié à Paris.")

print(f'   Avant: {para1.text}')
print(f'   Runs: {len(para1.runs)}')
for i, run in enumerate(para1.runs):
    print(f'      Run {i}: "{run.text}" (bold={run.bold})')

# Remplacer avec formatage
success = replace_placeholder_with_format(para1, "[NOM DU PRENEUR]", "Jean Dupont")
print(f'   Remplacement: {"✅ OK" if success else "❌ ÉCHEC"}')
print(f'   Après: {para1.text}')
print(f'   Runs: {len(para1.runs)}')
for i, run in enumerate(para1.runs):
    print(f'      Run {i}: "{run.text}" (bold={run.bold})')

# Test 2: Placeholder normal (pas de formatage)
print('\n2️⃣  Test placeholder normal')
para2 = doc.add_paragraph()
run1 = para2.add_run("Le montant est [MONTANT] euros.")

print(f'   Avant: {para2.text}')
print(f'   Runs: {len(para2.runs)}')

success = replace_placeholder_with_format(para2, "[MONTANT]", "5000")
print(f'   Remplacement: {"✅ OK" if success else "❌ ÉCHEC"}')
print(f'   Après: {para2.text}')

# Test 3: Placeholder en italique
print('\n3️⃣  Test placeholder en italique')
para3 = doc.add_paragraph()
run1 = para3.add_run("Note: ")
run2 = para3.add_run("[NOTE]")
run2.italic = True
run3 = para3.add_run(" est important.")

print(f'   Avant: {para3.text}')
success = replace_placeholder_with_format(para3, "[NOTE]", "Ceci")
print(f'   Remplacement: {"✅ OK" if success else "❌ ÉCHEC"}')
print(f'   Après: {para3.text}')
for i, run in enumerate(para3.runs):
    print(f'      Run {i}: "{run.text}" (italic={run.italic})')

# Sauvegarder pour inspection manuelle
doc.save('test_placeholder_formatting_output.docx')
print('\n💾 Document sauvegardé: test_placeholder_formatting_output.docx')
print('   Ouvrez-le dans Word pour vérifier le formatage')

print('\n' + '=' * 80)
print('TEST TERMINÉ')
print('=' * 80)
