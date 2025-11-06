"""Crée le template Word BAIL avec placeholders."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Créer un nouveau document
doc = Document()

# Configuration des marges
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Titre principal
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("BAIL COMMERCIAL")
title_run.bold = True
title_run.font.size = Pt(16)

doc.add_paragraph()  # Ligne vide

# Sous-titre
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("ENTRE LES SOUSSIGNES :")
subtitle_run.bold = True
subtitle_run.font.size = Pt(14)

doc.add_paragraph()  # Ligne vide

# Section Comparution Bailleur
doc.add_paragraph("{{COMPARUTION_BAILLEUR}}")
doc.add_paragraph()

partie1 = doc.add_paragraph()
partie1.alignment = WD_ALIGN_PARAGRAPH.CENTER
partie1_run = partie1.add_run("D'UNE PART,")
partie1_run.bold = True

doc.add_paragraph()
doc.add_paragraph("ET :")
doc.add_paragraph()

# Section Comparution Preneur
doc.add_paragraph("{{COMPARUTION_PRENEUR}}")
doc.add_paragraph()

partie2 = doc.add_paragraph()
partie2.alignment = WD_ALIGN_PARAGRAPH.CENTER
partie2_run = partie2.add_run("D'AUTRE PART,")
partie2_run.bold = True

doc.add_paragraph()

# IL A ETE CONVENU
convenu = doc.add_paragraph()
convenu.alignment = WD_ALIGN_PARAGRAPH.CENTER
convenu_run = convenu.add_run("IL A ETE CONVENU ET ARRETE CE QUI SUIT :")
convenu_run.bold = True

doc.add_paragraph()
doc.add_paragraph("_" * 80)
doc.add_paragraph()

# Article préliminaire (conditionnel)
preliminaire_header = doc.add_paragraph()
preliminaire_run = preliminaire_header.add_run("{{ARTICLE_PRELIMINAIRE}}")

doc.add_paragraph()
doc.add_paragraph("_" * 80)
doc.add_paragraph()

# ARTICLE 1 - DESIGNATION
art1_title = doc.add_paragraph()
art1_title_run = art1_title.add_run("ARTICLE 1 – DESIGNATION")
art1_title_run.bold = True
art1_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_1}}")
doc.add_paragraph()

# ARTICLE 2 - DUREE
art2_title = doc.add_paragraph()
art2_title_run = art2_title.add_run("ARTICLE 2 – DUREE")
art2_title_run.bold = True
art2_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_2}}")
doc.add_paragraph()

# ARTICLE 3 - DESTINATION
art3_title = doc.add_paragraph()
art3_title_run = art3_title.add_run("ARTICLE 3 – DESTINATION")
art3_title_run.bold = True
art3_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_3}}")
doc.add_paragraph()

# ARTICLE 4 - ETAT DES LIEUX
art4_title = doc.add_paragraph()
art4_title_run = art4_title.add_run("ARTICLE 4 – ETAT DES LIEUX")
art4_title_run.bold = True
art4_title_run.font.size = Pt(12)

doc.add_paragraph("[Standard - non géré par Excel actuellement]")
doc.add_paragraph()

# ARTICLE 5 - ENTRETIEN ET REPARATIONS
art5_title = doc.add_paragraph()
art5_title_run = art5_title.add_run("ARTICLE 5 – ENTRETIEN ET REPARATIONS")
art5_title_run.bold = True
art5_title_run.font.size = Pt(12)

doc.add_paragraph("5.3. {{ARTICLE_5_3}}")
doc.add_paragraph()

# ARTICLE 6 - TRAVAUX
art6_title = doc.add_paragraph()
art6_title_run = art6_title.add_run("ARTICLE 6 – TRAVAUX")
art6_title_run.bold = True
art6_title_run.font.size = Pt(12)

doc.add_paragraph("[Standard - non géré par Excel actuellement]")
doc.add_paragraph()

# ARTICLE 7 - LOYER
art7_title = doc.add_paragraph()
art7_title_run = art7_title.add_run("ARTICLE 7 – LOYER")
art7_title_run.bold = True
art7_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_7_1}}")
doc.add_paragraph()
doc.add_paragraph("{{ARTICLE_7_2}}")
doc.add_paragraph()
doc.add_paragraph("{{ARTICLE_7_3}}")
doc.add_paragraph()
doc.add_paragraph("{{ARTICLE_7_6}}")
doc.add_paragraph()

# ARTICLE 8 - GARANTIES
art8_title = doc.add_paragraph()
art8_title_run = art8_title.add_run("ARTICLE 8 – GARANTIES")
art8_title_run.bold = True
art8_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_8}}")
doc.add_paragraph()

# ARTICLE 19 - FRAIS ET HONORAIRES
art19_title = doc.add_paragraph()
art19_title_run = art19_title.add_run("ARTICLE 19 – FRAIS ET HONORAIRES")
art19_title_run.bold = True
art19_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_19}}")
doc.add_paragraph()

# ARTICLE 22.2 - DPE
art22_title = doc.add_paragraph()
art22_title_run = art22_title.add_run("ARTICLE 22.2 – DIAGNOSTIC DE PERFORMANCE ENERGETIQUE")
art22_title_run.bold = True
art22_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_22_2}}")
doc.add_paragraph()

# ARTICLE 26 - DISPOSITIONS PARTICULIERES
art26_title = doc.add_paragraph()
art26_title_run = art26_title.add_run("ARTICLE 26 – DISPOSITIONS PARTICULIERES")
art26_title_run.bold = True
art26_title_run.font.size = Pt(12)

doc.add_paragraph("{{ARTICLE_26}}")
doc.add_paragraph()
doc.add_paragraph("{{ARTICLE_26_1}}")
doc.add_paragraph()
doc.add_paragraph("{{ARTICLE_26_2}}")
doc.add_paragraph()

# Signature
doc.add_paragraph()
doc.add_paragraph("_" * 80)
doc.add_paragraph()

fait = doc.add_paragraph()
fait_run = fait.add_run("Fait à {{VILLE}}, le {{DATE_SIGNATURE}}")
fait_run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Signatures
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

for cell in table.rows[0].cells:
    cell.width = Inches(3)

cell_bailleur = table.rows[0].cells[0]
cell_preneur = table.rows[0].cells[1]

p_bailleur = cell_bailleur.paragraphs[0]
p_bailleur.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_bailleur_run = p_bailleur.add_run("Le Bailleur")
p_bailleur_run.bold = True

p_preneur = cell_preneur.paragraphs[0]
p_preneur.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_preneur_run = p_preneur.add_run("Le Preneur")
p_preneur_run.bold = True

# Sauvegarder
output_path = "Template BAIL avec placeholder.docx"
doc.save(output_path)

print("=" * 80)
print("✅ TEMPLATE BAIL CRÉÉ AVEC SUCCÈS")
print("=" * 80)
print(f"\nFichier: {output_path}")
print("\nPlaceholders utilisés:")
placeholders = [
    "{{COMPARUTION_BAILLEUR}}",
    "{{COMPARUTION_PRENEUR}}",
    "{{ARTICLE_PRELIMINAIRE}}",
    "{{ARTICLE_1}}",
    "{{ARTICLE_2}}",
    "{{ARTICLE_3}}",
    "{{ARTICLE_5_3}}",
    "{{ARTICLE_7_1}}",
    "{{ARTICLE_7_2}}",
    "{{ARTICLE_7_3}}",
    "{{ARTICLE_7_6}}",
    "{{ARTICLE_8}}",
    "{{ARTICLE_19}}",
    "{{ARTICLE_22_2}}",
    "{{ARTICLE_26}}",
    "{{ARTICLE_26_1}}",
    "{{ARTICLE_26_2}}",
    "{{VILLE}}",
    "{{DATE_SIGNATURE}}"
]

for i, ph in enumerate(placeholders, 1):
    print(f"  {i}. {ph}")

print("\n" + "=" * 80)
