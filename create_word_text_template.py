"""
Création d'un document Word contenant tous les textes du fichier Excel BAIL.
L'utilisateur pourra ensuite formater ce document (gras, italique, titres, etc.)
et le système chargera ces styles formatés.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from pathlib import Path


def create_word_text_template():
    """Crée un document Word avec tous les textes du fichier Excel BAIL."""

    print('=' * 80)
    print('CRÉATION DU DOCUMENT WORD DE TEXTES POUR FORMATAGE')
    print('=' * 80)

    # Charger le fichier Excel
    print('\n📄 Chargement du fichier Excel...')
    excel_path = Path('Redaction BAIL.xlsx')

    if not excel_path.exists():
        print(f'❌ Fichier non trouvé: {excel_path}')
        return None

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb['Rédaction BAIL']

    print(f'   ✅ Fichier chargé ({ws.max_row} lignes)')

    # Créer le document Word
    print('\n📝 Création du document Word...')
    doc = Document()

    # Ajouter un titre principal
    title = doc.add_heading('TEXTES BAIL - À FORMATER', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter instructions
    instructions = doc.add_paragraph()
    instructions.add_run('INSTRUCTIONS:\n').bold = True
    instructions.add_run(
        'Ce document contient tous les textes qui seront utilisés pour générer les baux.\n'
        'Vous pouvez maintenant formater ces textes comme vous le souhaitez:\n'
        '  • Mettre du texte en gras, italique, souligné\n'
        '  • Changer la police, la taille, la couleur\n'
        '  • Appliquer des styles de titre (Titre 1, Titre 2, etc.)\n\n'
        'Les formats seront automatiquement détectés et appliqués dans les documents générés.\n\n'
        'NE PAS modifier les identifiants en MAJUSCULES (ex: "ARTICLE_1") - '
        'ils servent à identifier le contenu.\n'
    )
    instructions.runs[-1].font.color.rgb = RGBColor(200, 50, 50)

    doc.add_page_break()

    # Extraire les en-têtes
    headers = {}
    for col in range(1, ws.max_column + 1):
        header = ws.cell(1, col).value
        if header:
            headers[col] = header

    # Identifier les colonnes de texte à extraire
    text_columns = {
        7: 'Entrée correspondante - Option 1',
        10: 'Entrée correspondante - Option 2',
    }

    # Dictionnaire pour stocker les textes par section
    sections = {}

    print('\n🔍 Extraction des textes...')

    # Garder une trace de l'article/désignation courant pour les lignes None
    current_article = None
    current_designation = None
    section_counter = {}  # Pour numéroter les sections multiples

    for row_idx in range(2, ws.max_row + 1):
        # Récupérer les métadonnées de la ligne
        article = ws.cell(row_idx, 1).value
        designation = ws.cell(row_idx, 2).value
        condition_opt1 = ws.cell(row_idx, 6).value
        condition_opt2 = ws.cell(row_idx, 9).value

        # Mettre à jour l'article/désignation courant si présent
        if article:
            current_article = article
            current_designation = designation

        # Utiliser l'article/désignation courant si la ligne est None
        effective_article = article if article else current_article
        effective_designation = designation if designation else current_designation

        # Si toujours aucun article, utiliser le numéro de ligne
        if not effective_article:
            effective_article = f"ROW_{row_idx}"

        # Créer un identifiant pour cette ligne
        section_id = str(effective_article).strip().upper().replace(' ', '_').replace('.', '_')
        if effective_designation:
            # Nettoyer la désignation pour l'ID
            clean_designation = str(effective_designation).strip().upper()
            clean_designation = clean_designation.replace(' ', '_')
            clean_designation = clean_designation.replace('/', '_')
            clean_designation = clean_designation.replace('-', '_')
            clean_designation = clean_designation.replace('.', '_')
            section_id = f"{section_id}_{clean_designation}"

        # Compter les occurrences pour différencier les lignes multiples
        base_section_id = section_id
        if base_section_id not in section_counter:
            section_counter[base_section_id] = 0
        else:
            section_counter[base_section_id] += 1
            section_id = f"{base_section_id}_{section_counter[base_section_id]}"

        # Extraire le texte Option 1
        text_opt1 = ws.cell(row_idx, 7).value
        if text_opt1:
            text_opt1 = str(text_opt1).strip()
            if text_opt1:
                key = f"{section_id}_OPTION_1" if condition_opt1 else section_id
                sections[key] = {
                    'article': effective_article,
                    'designation': effective_designation,
                    'condition': condition_opt1,
                    'text': text_opt1,
                    'row': row_idx
                }

        # Extraire le texte Option 2
        text_opt2 = ws.cell(row_idx, 10).value
        if text_opt2:
            text_opt2 = str(text_opt2).strip()
            if text_opt2:
                key = f"{section_id}_OPTION_2"
                sections[key] = {
                    'article': effective_article,
                    'designation': effective_designation,
                    'condition': condition_opt2,
                    'text': text_opt2,
                    'row': row_idx
                }

    print(f'   ✅ {len(sections)} sections de texte extraites')

    # Ajouter les sections au document
    print('\n📋 Ajout des sections au document...')

    current_article = None
    section_count = 0

    for section_id, section_data in sorted(sections.items(), key=lambda x: x[1]['row']):
        article = section_data['article']
        designation = section_data['designation']
        condition = section_data['condition']
        text = section_data['text']

        # Ajouter séparateur d'article si changement
        if article != current_article:
            if current_article is not None:
                doc.add_page_break()

            # Titre de l'article
            article_heading = doc.add_heading(f'ARTICLE: {article}', level=1)
            article_heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
            current_article = article

        # Ajouter l'identifiant (en commentaire visuel)
        id_para = doc.add_paragraph()
        id_run = id_para.add_run(f'[ID: {section_id}]')
        id_run.font.size = Pt(9)
        id_run.font.color.rgb = RGBColor(150, 150, 150)
        id_run.italic = True

        # Ajouter la désignation si présente
        if designation:
            designation_para = doc.add_paragraph()
            designation_run = designation_para.add_run(f'Désignation: {designation}')
            designation_run.font.size = Pt(10)
            designation_run.bold = True
            designation_run.font.color.rgb = RGBColor(0, 100, 0)

        # Ajouter la condition si présente
        if condition:
            condition_para = doc.add_paragraph()
            condition_run = condition_para.add_run(f'Condition: {condition}')
            condition_run.font.size = Pt(9)
            condition_run.italic = True
            condition_run.font.color.rgb = RGBColor(180, 100, 0)

        # Ajouter le texte (c'est ici que l'utilisateur va formater)
        text_para = doc.add_paragraph()
        text_para.add_run(text)

        # Ajouter un séparateur
        doc.add_paragraph('─' * 80)

        section_count += 1

    print(f'   ✅ {section_count} sections ajoutées')

    # Sauvegarder le document
    output_path = Path('Textes BAIL avec styles.docx')
    print(f'\n💾 Sauvegarde du document...')
    print(f'   Fichier: {output_path}')

    doc.save(output_path)

    print(f'   ✅ Document sauvegardé ({len(doc.paragraphs)} paragraphes)')

    # Résumé
    print('\n' + '=' * 80)
    print('RÉSUMÉ')
    print('=' * 80)
    print(f'✅ Document créé: {output_path}')
    print(f'✅ Sections de texte: {section_count}')
    print(f'✅ Articles couverts: {len(set(s["article"] for s in sections.values()))}')
    print()
    print('📝 PROCHAINES ÉTAPES:')
    print('   1. Ouvrir le document dans Word')
    print('   2. Formater les textes comme souhaité (gras, italique, titres, etc.)')
    print('   3. Sauvegarder le document')
    print('   4. Le système chargera automatiquement les styles formatés')
    print()
    print('⚠️  IMPORTANT:')
    print('   - Ne pas modifier les identifiants [ID: ...]')
    print('   - Ne pas supprimer de sections')
    print('   - Vous pouvez formater librement le texte de chaque section')

    return output_path


if __name__ == "__main__":
    result = create_word_text_template()
    if result:
        print(f'\n🎉 SUCCESS! Document créé: {result}')
    else:
        print('\n❌ ÉCHEC de la création du document')
