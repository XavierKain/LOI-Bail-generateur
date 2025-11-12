"""Add president names to signature table in template"""

from docx import Document
from pathlib import Path

template_path = "Template BAIL avec placeholder.docx"

print("=" * 80)
print("AJOUT DES NOMS DE PRÉSIDENTS AU TABLEAU DE SIGNATURES")
print("=" * 80)

# Charger le template
doc = Document(template_path)

# Trouver le tableau de signatures (premier tableau)
signature_table = doc.tables[0]

print(f"\n📋 Tableau actuel:")
print(f"  Lignes: {len(signature_table.rows)}")
print(f"  Colonnes: {len(signature_table.columns)}")

# Vérifier la structure actuelle
row0 = signature_table.rows[0]
cell0_text = row0.cells[0].text.strip()
cell1_text = row0.cells[1].text.strip()

print(f"\n  Ligne 0:")
print(f"    Col 0: '{cell0_text}'")
print(f"    Col 1: '{cell1_text}'")

# Ajouter une nouvelle ligne pour les noms
new_row = signature_table.add_row()

# Colonne 0: Nom du président Bailleur (toujours Maxime FORGEOT)
new_row.cells[0].text = "Monsieur Maxime FORGEOT"

# Colonne 1: Placeholder pour le président Preneur (sera remplacé par INPI)
new_row.cells[1].text = "[PRESIDENT DE LA SOCIETE]"

print(f"\n✅ Nouvelle ligne ajoutée:")
print(f"  Ligne 1:")
print(f"    Col 0: '{new_row.cells[0].text}'")
print(f"    Col 1: '{new_row.cells[1].text}'")

# Créer une backup du template original
backup_path = Path(template_path).with_suffix('.backup.docx')
if not backup_path.exists():
    import shutil
    shutil.copy2(template_path, backup_path)
    print(f"\n💾 Backup créé: {backup_path}")

# Sauvegarder le template modifié
doc.save(template_path)
print(f"\n✅ Template modifié sauvegardé: {template_path}")

print("\n" + "=" * 80)
print("VERIFICATION:")
print("=" * 80)

# Recharger pour vérifier
doc_verify = Document(template_path)
table_verify = doc_verify.tables[0]

print(f"\nTableau après modification:")
for row_idx, row in enumerate(table_verify.rows):
    print(f"\n  Ligne {row_idx}:")
    for cell_idx, cell in enumerate(row.cells):
        print(f"    Col {cell_idx}: '{cell.text.strip()}'")

print("\n" + "=" * 80)
