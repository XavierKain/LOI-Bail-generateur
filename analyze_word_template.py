"""Analyse du template Word BAIL pour comprendre la structure."""

from docx import Document
import os

template_path = "2024 - Bail type.doc"

if not os.path.exists(template_path):
    print(f"❌ Fichier non trouvé: {template_path}")
    exit(1)

try:
    # Ouvrir le document
    doc = Document(template_path)

    print("=" * 80)
    print("ANALYSE DU TEMPLATE BAIL")
    print("=" * 80)

    print(f"\nNombre de paragraphes: {len(doc.paragraphs)}")
    print(f"Nombre de sections: {len(doc.sections)}")

    # Analyser les styles utilisés
    styles_used = set()
    for para in doc.paragraphs:
        if para.style:
            styles_used.add(para.style.name)

    print(f"\nStyles utilisés ({len(styles_used)}):")
    for style in sorted(styles_used):
        print(f"  - {style}")

    # Afficher la structure (premiers paragraphes)
    print(f"\n{'=' * 80}")
    print("STRUCTURE DU DOCUMENT (50 premiers paragraphes)")
    print('=' * 80)

    for i, para in enumerate(doc.paragraphs[:50], 1):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "None"
            preview = text[:80] if len(text) > 80 else text
            print(f"\n{i}. [{style}]")
            print(f"   {preview}...")

    # Chercher les titres d'articles
    print(f"\n{'=' * 80}")
    print("ARTICLES DÉTECTÉS")
    print('=' * 80)

    articles_found = []
    for para in doc.paragraphs:
        text = para.text.strip().upper()
        if text.startswith("ARTICLE"):
            articles_found.append(text)

    for i, article in enumerate(articles_found, 1):
        print(f"{i}. {article}")

    print(f"\n{'=' * 80}")
    print(f"Total: {len(articles_found)} articles détectés")
    print('=' * 80)

except Exception as e:
    print(f"❌ Erreur lors de l'analyse: {e}")
    import traceback
    traceback.print_exc()
