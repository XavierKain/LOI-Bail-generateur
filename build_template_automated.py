"""
Construction automatisée du template BAIL complet.
Copie les articles depuis le document original Word en préservant la mise en forme.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from copy import deepcopy

def load_original_doc():
    """Charge le document original converti en .docx."""
    print("📄 Chargement du document original...")
    try:
        # Essayer d'abord la version convertie
        doc = Document("2024 - Bail type CONVERTED.docx")
        print("   ✅ Document converti chargé")
        return doc
    except Exception as e:
        print(f"   ❌ Erreur: {e}")
        return None

def find_article_in_doc(doc, article_id):
    """
    Trouve un article dans le document et retourne ses paragraphes.

    Args:
        doc: Document source
        article_id: 'PRELIMINAIRE' ou numéro (9, 10, etc.)

    Returns:
        Liste de paragraphes (objets paragraph) de l'article
    """
    if article_id == 'PRELIMINAIRE':
        start_pattern = r'^ARTICLE PRELIMINAIRE'
        next_pattern = r'^ARTICLE\s+1\s'
    else:
        start_pattern = f'^ARTICLE\\s+{article_id}\\s'
        next_pattern = f'^ARTICLE\\s+{article_id + 1}\\s'

    start_idx = None
    end_idx = len(doc.paragraphs)

    # Trouver le début
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(start_pattern, text, re.IGNORECASE):
            start_idx = i
            break

    if start_idx is None:
        return None

    # Trouver la fin (prochain article)
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if re.match(r'^ARTICLE\s+\d+', text, re.IGNORECASE):
            end_idx = i
            break

    return doc.paragraphs[start_idx:end_idx]

def copy_paragraph_with_format(source_para):
    """
    Copie un paragraphe avec son formatage.
    Retourne un dictionnaire avec les propriétés du paragraphe.
    """
    para_data = {
        'text': source_para.text,
        'runs': [],
        'style': source_para.style.name if source_para.style else None,
        'alignment': source_para.alignment,
    }

    # Copier chaque run avec son formatage
    for run in source_para.runs:
        run_data = {
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
        }
        para_data['runs'].append(run_data)

    return para_data

def add_paragraph_with_format(doc, para_data):
    """Ajoute un paragraphe formaté au document."""
    para = doc.add_paragraph()

    # Appliquer le style
    if para_data['style']:
        try:
            para.style = para_data['style']
        except:
            pass

    # Appliquer l'alignement
    if para_data['alignment']:
        para.alignment = para_data['alignment']

    # Ajouter les runs avec formatage
    if para_data['runs']:
        para.clear()  # Supprimer le contenu par défaut
        for run_data in para_data['runs']:
            run = para.add_run(run_data['text'])

            if run_data['bold'] is not None:
                run.bold = run_data['bold']
            if run_data['italic'] is not None:
                run.italic = run_data['italic']
            if run_data['underline'] is not None:
                run.underline = run_data['underline']
            if run_data['font_name']:
                run.font.name = run_data['font_name']
            if run_data['font_size']:
                run.font.size = run_data['font_size']
            if run_data['font_color']:
                run.font.color.rgb = run_data['font_color']
    else:
        # Fallback: utiliser le texte brut
        para.text = para_data['text']

    return para

def extract_existing_articles(template_doc):
    """Extrait les articles existants du template actuel."""
    print("📋 Extraction des articles existants du template...")

    existing = {}
    current_article = None
    article_paras = []

    for para in template_doc.paragraphs:
        text = para.text.strip()
        match = re.match(r'^ARTICLE\s+(\d+)', text, re.IGNORECASE)

        if match:
            # Sauvegarder l'article précédent
            if current_article is not None and article_paras:
                existing[current_article] = article_paras

            # Nouveau article
            current_article = int(match.group(1))
            article_paras = [copy_paragraph_with_format(para)]
        elif current_article is not None:
            article_paras.append(copy_paragraph_with_format(para))

    # Sauvegarder le dernier
    if current_article is not None and article_paras:
        existing[current_article] = article_paras

    print(f"   ✅ {len(existing)} articles existants extraits: {sorted(existing.keys())}")
    return existing

def build_complete_template():
    """Construit le template complet automatiquement."""

    print('=' * 80)
    print('CONSTRUCTION AUTOMATISÉE DU TEMPLATE BAIL COMPLET')
    print('=' * 80)

    # Charger les documents
    print('\n1️⃣  CHARGEMENT DES DOCUMENTS')
    print('-' * 80)

    original_doc = load_original_doc()
    if not original_doc:
        print("❌ Impossible de charger le document original")
        return None

    template_doc = Document("Template BAIL avec placeholder.docx")
    print(f"   ✅ Template actuel chargé ({len(template_doc.paragraphs)} paragraphes)")

    # Extraire les articles existants
    existing_articles = extract_existing_articles(template_doc)

    # Créer nouveau document
    print('\n2️⃣  CRÉATION DU NOUVEAU TEMPLATE')
    print('-' * 80)

    new_doc = Document()

    # Copier les paramètres de section
    for section in template_doc.sections:
        new_section = new_doc.sections[0] if len(new_doc.sections) == 1 else new_doc.add_section()
        new_section.page_height = section.page_height
        new_section.page_width = section.page_width
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin

    # Copier l'en-tête (parties)
    print("\n   📝 Copie de l'en-tête (parties)...")
    header_copied = False
    for para in template_doc.paragraphs:
        if para.text.strip().startswith('ARTICLE'):
            break
        para_data = copy_paragraph_with_format(para)
        add_paragraph_with_format(new_doc, para_data)
        header_copied = True

    if header_copied:
        print('   ✅ En-tête copié')

    # Ajouter page break pour TOC
    new_doc.add_page_break()

    # Ajouter titre TOC
    print('\n   📑 Ajout de la table des matières...')
    toc_title = new_doc.add_paragraph("TABLE DES MATIÈRES")
    toc_title.runs[0].bold = True
    toc_title.runs[0].font.size = Pt(16)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note: TOC dynamique doit être ajoutée manuellement dans Word
    # ou via python-docx-template (plus complexe)
    toc_note = new_doc.add_paragraph()
    toc_note.add_run("[La table des matières sera générée automatiquement par Word]").italic = True
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_page_break()

    # Construire la liste complète des articles
    print('\n3️⃣  AJOUT DES ARTICLES')
    print('-' * 80)

    articles_to_add = ['PRELIMINAIRE'] + list(range(1, 29))

    articles_added = 0
    articles_from_template = 0
    articles_from_original = 0

    for art_id in articles_to_add:
        art_int = art_id if art_id == 'PRELIMINAIRE' else int(art_id)

        # Priorité au template existant
        if art_int != 'PRELIMINAIRE' and art_int in existing_articles:
            print(f'   ✅ ARTICLE {art_int} (depuis template existant)')
            for para_data in existing_articles[art_int]:
                add_paragraph_with_format(new_doc, para_data)
            articles_added += 1
            articles_from_template += 1
        else:
            # Essayer de copier depuis l'original
            # Note: Le document converti perd la mise en forme, on va copier le texte
            # et appliquer un formatage basique

            if art_id == 'PRELIMINAIRE':
                article_file = 'article_PRELIMINAIRE_extracted.txt'
            else:
                article_file = f'article_{art_id}_extracted.txt'

            try:
                with open(article_file, 'r', encoding='utf-8') as f:
                    content = f.read()

                lines = content.split('\n')
                for i, line in enumerate(lines):
                    if not line.strip():
                        new_doc.add_paragraph()  # Ligne vide
                        continue

                    para = new_doc.add_paragraph(line)

                    # Appliquer formatage basique
                    # Titre d'article en gras
                    if i == 0 and line.strip().startswith('ARTICLE'):
                        para.runs[0].bold = True
                        para.runs[0].font.size = Pt(12)
                    # Sous-sections en gras
                    elif re.match(r'^\d+\.\d+\.', line.strip()):
                        para.runs[0].bold = True

                print(f'   ✅ ARTICLE {art_id} (depuis fichier extrait, {len(lines)} lignes)')
                articles_added += 1
                articles_from_original += 1

            except FileNotFoundError:
                print(f'   ⚠️  ARTICLE {art_id} non trouvé (fichier manquant)')

    print(f'\n   📊 Total: {articles_added} articles ajoutés')
    print(f'      - Depuis template existant: {articles_from_template}')
    print(f'      - Depuis original: {articles_from_original}')

    # Ajouter numérotation des pages
    print('\n4️⃣  AJOUT DE LA NUMÉROTATION DES PAGES')
    print('-' * 80)

    for i, section in enumerate(new_doc.sections):
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()

        # Ajouter champ PAGE
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.size = Pt(10)

        print(f'   ✅ Section {i+1}: numérotation ajoutée')

    # Sauvegarder
    output_file = "Template BAIL avec placeholder AUTOMATISE.docx"
    print(f'\n5️⃣  SAUVEGARDE')
    print('-' * 80)
    print(f'   💾 Fichier: {output_file}')

    new_doc.save(output_file)

    print(f'   ✅ Template sauvegardé ({len(new_doc.paragraphs)} paragraphes)')

    # Résumé
    print('\n' + '=' * 80)
    print('RÉSUMÉ')
    print('=' * 80)
    print(f'✅ Template créé: {output_file}')
    print(f'✅ Articles: {articles_added}/29')
    print(f'✅ Table des matières: Placeholder ajouté')
    print(f'✅ Numérotation: Activée')
    print('\n⚠️  ÉTAPE MANUELLE REQUISE:')
    print('   Ouvrir le document dans Word et:')
    print('   1. Remplacer le placeholder TOC par une vraie table Word')
    print('   2. Appliquer les styles Titre 2/3 aux articles')
    print('   3. Mettre à jour la TOC (clic droit → Mettre à jour)')

    return output_file

if __name__ == "__main__":
    output = build_complete_template()
    if output:
        print(f'\n🎉 SUCCESS! Template créé: {output}')
    else:
        print('\n❌ ÉCHEC de la création du template')
