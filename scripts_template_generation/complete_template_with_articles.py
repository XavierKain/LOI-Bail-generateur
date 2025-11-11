"""Complète le template BAIL actuel avec tous les articles manquants."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def find_article_position(doc, article_num):
    """Trouve la position (index de paragraphe) où insérer un article."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Chercher ARTICLE X
        match = re.match(r'^ARTICLE\s+(\d+)', text)
        if match:
            found_num = int(match.group(1))
            if found_num > article_num:
                # Insérer avant cet article
                return i
            elif found_num == article_num:
                # Article déjà présent
                return None

    # Si pas trouvé, insérer à la fin
    return len(doc.paragraphs)

def insert_article_at_position(doc, article_content, insert_before_para_idx):
    """Insère un article à une position donnée."""
    # Séparer le contenu en lignes
    lines = article_content.split('\n')

    # Insérer chaque ligne comme paragraphe
    for line_idx, line in enumerate(lines):
        # Créer un nouveau paragraphe à la position
        p = doc.paragraphs[insert_before_para_idx]._element
        new_p = p.getparent().insert(insert_before_para_idx, p.__class__())

        # Ajouter le texte
        new_para = doc.paragraphs[insert_before_para_idx]
        new_para.text = line

        # Style pour le titre d'article
        if line_idx == 0 and line.strip().startswith('ARTICLE'):
            new_para.runs[0].bold = True
            new_para.runs[0].font.size = Pt(12)

        insert_before_para_idx += 1

    return insert_before_para_idx

def complete_template():
    """Complète le template avec tous les articles manquants."""

    print('=' * 80)
    print('COMPLÉTION DU TEMPLATE BAIL')
    print('=' * 80)

    # Charger le template actuel
    template_file = "Template BAIL avec placeholder.docx"
    print(f'\n📄 Chargement: {template_file}')
    doc = Document(template_file)
    print(f'   Paragraphes initiaux: {len(doc.paragraphs)}')

    # Articles à insérer (dans l'ordre décroissant pour éviter les décalages d'index)
    articles_to_insert = [
        ('PRELIMINAIRE', 0),  # Insérer au début
        (9, 9),
        (10, 10),
        (11, 11),
        (12, 12),
        (13, 13),
        (14, 14),
        (15, 15),
        (16, 16),
        (17, 17),
        (18, 18),
        (20, 20),
        (21, 21),
        (23, 23),
        (24, 24),
        (25, 25),
        (27, 27),
        (28, 28),
    ]

    # Inverser l'ordre pour insérer du plus grand au plus petit numéro
    # (évite les problèmes de décalage d'index)
    articles_to_insert.reverse()

    inserted_count = 0

    for art_id, art_num in articles_to_insert:
        article_file = f'article_{art_id}_extracted.txt'

        try:
            # Charger le contenu de l'article
            with open(article_file, 'r', encoding='utf-8') as f:
                content = f.read().strip()

            print(f'\n📝 Insertion ARTICLE {art_id}')

            # Trouver la position d'insertion
            if art_id == 'PRELIMINAIRE':
                # Insérer au tout début (après le titre)
                position = 3  # Après titre, parties, etc.
            else:
                position = find_article_position(doc, art_num)

            if position is None:
                print(f'   ⚠️  Article {art_id} déjà présent, skip')
                continue

            print(f'   Position: paragraphe {position}')

            # Insérer l'article
            # Note: on utilise une méthode plus simple - ajouter à la fin puis réorganiser
            # Pour simplifier, on va recréer le document dans le bon ordre

            inserted_count += 1

        except FileNotFoundError:
            print(f'   ❌ Fichier {article_file} non trouvé')
            continue

    print(f'\n✅ {inserted_count} articles à insérer')

    # Approche alternative: reconstruire le document dans le bon ordre
    print('\n🔄 Reconstruction du document dans l\'ordre correct...')

    # Créer un nouveau document
    new_doc = Document()

    # Copier les paramètres de section
    for section in doc.sections:
        new_section = new_doc.sections[0] if len(new_doc.sections) > 0 else new_doc.sections[-1]
        new_section.page_height = section.page_height
        new_section.page_width = section.page_width
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin

    # Récupérer tous les articles du template actuel
    existing_articles = {}
    current_article = None
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r'^ARTICLE\s+(\d+)', text)

        if match:
            # Sauvegarder l'article précédent
            if current_article is not None:
                existing_articles[current_article] = '\n'.join(current_content)

            # Nouveau article
            current_article = int(match.group(1))
            current_content = [text]
        elif current_article is not None:
            current_content.append(text)

    # Sauvegarder le dernier article
    if current_article is not None:
        existing_articles[current_article] = '\n'.join(current_content)

    print(f'   Articles existants: {sorted(existing_articles.keys())}')

    # Charger tous les articles extraits
    extracted_articles = {}
    for art_id, _ in articles_to_insert:
        article_file = f'article_{art_id}_extracted.txt'
        try:
            with open(article_file, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                extracted_articles[art_id] = content
        except FileNotFoundError:
            pass

    print(f'   Articles extraits: {sorted([str(k) for k in extracted_articles.keys()])}')

    # Reconstruire dans l'ordre: PRELIMINAIRE, 1-28
    all_article_nums = ['PRELIMINAIRE'] + list(range(1, 29))

    # Copier le début du document (avant Article 1)
    for para in doc.paragraphs:
        if para.text.strip().startswith('ARTICLE'):
            break
        new_para = new_doc.add_paragraph(para.text)
        # Copier le style
        if para.style:
            new_para.style = para.style

    # Ajouter tous les articles dans l'ordre
    for art_num in all_article_nums:
        # Vérifier si article existe (template actuel ou extrait)
        if art_num == 'PRELIMINAIRE':
            if art_num in extracted_articles:
                print(f'   ✅ Ajout ARTICLE PRELIMINAIRE (extrait)')
                for line in extracted_articles[art_num].split('\n'):
                    new_doc.add_paragraph(line)
            continue

        art_int = int(art_num) if art_num != 'PRELIMINAIRE' else 0

        # Priorité au template actuel
        if art_int in existing_articles:
            print(f'   ✅ Ajout ARTICLE {art_int} (template actuel)')
            for line in existing_articles[art_int].split('\n'):
                new_doc.add_paragraph(line)
        elif art_int in extracted_articles:
            print(f'   ✅ Ajout ARTICLE {art_int} (extrait)')
            for line in extracted_articles[art_int].split('\n'):
                new_doc.add_paragraph(line)

    # Sauvegarder le nouveau template
    output_file = "Template BAIL avec placeholder COMPLET.docx"
    new_doc.save(output_file)

    print(f'\n💾 Template complet sauvegardé: {output_file}')
    print(f'   Paragraphes finaux: {len(new_doc.paragraphs)}')

    return output_file

if __name__ == "__main__":
    output = complete_template()

    print('\n' + '=' * 80)
    print('RÉSUMÉ')
    print('=' * 80)
    print(f'✅ Template complet créé: {output}')
    print('\n📋 Prochaines étapes:')
    print('   1. Vérifier le template')
    print('   2. Ajouter la table des matières')
    print('   3. Ajouter la numérotation des pages')
