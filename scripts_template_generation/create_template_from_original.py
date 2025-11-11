"""Créer le template BAIL avec placeholders à partir du document original converti."""

from docx import Document
import re
from copy import deepcopy

def find_fields_to_replace(doc):
    """Identifie tous les champs à remplacer par des placeholders."""

    # Mapping des patterns à remplacer vers les placeholders
    replacements = {
        # Format [.] ou [●] -> placeholders spécifiques selon le contexte
        r'\[\.+\]': 'PLACEHOLDER',
        r'\[●\]': 'PLACEHOLDER',
        r'\[PRENOMS \+ NOM\]': '[PRESIDENT DE LA SOCIETE]',
        r'\[NOM DU NOTAIRE\]': '[NOM DU NOTAIRE]',
        r'\[DATE DU MARIAGE\]': '[DATE DU MARIAGE]',
        r'XXXXXXXXXX': '[NOM DU PRENEUR]',
        r'XXXXXXX': '[PLACEHOLDER]',
    }

    found_patterns = {}

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text

        # Chercher les patterns
        for pattern in [r'\[\.+\]', r'\[●\]', r'XXX+', r'\[.*?\]']:
            matches = re.finditer(pattern, text)
            for match in matches:
                key = match.group()
                if key not in found_patterns:
                    found_patterns[key] = []
                found_patterns[key].append((para_idx, text[:100]))

    return found_patterns

def create_template_with_placeholders(input_docx, output_docx):
    """Crée le template en remplaçant les champs par des placeholders."""

    print('=' * 80)
    print('CRÉATION DU TEMPLATE AVEC PLACEHOLDERS')
    print('=' * 80)

    # Charger le document converti
    doc = Document(input_docx)

    print(f'\n📄 Chargement: {input_docx}')
    print(f'   Nombre de paragraphes: {len(doc.paragraphs)}')
    print(f'   Nombre de sections: {len(doc.sections)}')

    # Analyser les champs à remplacer
    print('\n🔍 Recherche des champs à remplacer...')
    patterns = find_fields_to_replace(doc)

    print(f'\n   Patterns trouvés: {len(patterns)}')
    for pattern, occurrences in sorted(patterns.items())[:20]:  # Premiers 20
        print(f'   - "{pattern}": {len(occurrences)} occurrences')

    # Définir les remplacements contextuels
    contextual_replacements = [
        # Identification des parties
        {
            'search': 'La société [.], Société [.] au capital de [.]',
            'replace': 'La société [NOM DU BAILLEUR], [TYPE DE SOCIETE BAILLEUR] au capital de [CAPITAL SOCIAL BAILLEUR]',
            'context': 'bailleur'
        },
        {
            'search': 'immatriculée au Registre du Commerce et des Sociétés de [.] sous le numéro [.]',
            'replace': 'immatriculée au Registre du Commerce et des Sociétés de [RCS BAILLEUR] sous le numéro [SIREN BAILLEUR]',
            'context': 'bailleur'
        },
        {
            'search': 'dont le siège social est à [.]',
            'replace': 'dont le siège social est à [ADRESSE BAILLEUR]',
            'context': 'bailleur'
        },
        # Preneur (similaire)
        {
            'search': 'La société [.], Société [.] au capital de [.], immatriculée',
            'replace': 'La société [NOM DU PRENEUR], [TYPE DE SOCIETE] au capital de [CAPITAL SOCIAL], immatriculée',
            'context': 'preneur'
        },
        # Locaux
        {
            'search': 'sis à [.]',
            'replace': 'sis à [Adresse du local]',
            'context': 'local'
        },
        # Durée
        {
            'search': 'pour une durée de [.] années',
            'replace': 'pour une durée de [Durée initiale du bail] années',
            'context': 'duree'
        },
        # Loyer
        {
            'search': 'moyennant un loyer annuel de [.]',
            'replace': 'moyennant un loyer annuel de [LOYER ANNUEL HT]',
            'context': 'loyer'
        },
    ]

    print('\n✏️  Application des remplacements...')
    replacements_made = 0

    # Parcourir tous les paragraphes
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        modified_text = original_text

        # Remplacements simples pattern par pattern
        # [.] -> selon contexte
        if '[.]' in modified_text:
            # Contexte Bailleur (début du document)
            if para_idx < 20:
                if 'Société' in modified_text and 'capital' in modified_text:
                    modified_text = modified_text.replace(
                        'La société [.], Société [.] au capital de [.]',
                        'La société [NOM DU BAILLEUR], [TYPE DE SOCIETE BAILLEUR] au capital de [CAPITAL SOCIAL BAILLEUR]'
                    )
                if 'Registre du Commerce' in modified_text:
                    modified_text = re.sub(
                        r'de \[\.+\] sous le numéro \[\.+\]',
                        'de [RCS BAILLEUR] sous le numéro [SIREN BAILLEUR]',
                        modified_text
                    )
                if 'siège social' in modified_text:
                    modified_text = re.sub(
                        r'à \[\.+\]',
                        'à [ADRESSE BAILLEUR]',
                        modified_text
                    )

            # Contexte Preneur (après "D'AUTRE PART")
            if para_idx >= 20 and para_idx < 50:
                if 'Société' in modified_text and 'capital' in modified_text:
                    modified_text = modified_text.replace(
                        'La société [.]',
                        'La société [NOM DU PRENEUR]'
                    )
                    modified_text = re.sub(
                        r'Société \[\.+\] au capital de \[\.+\]',
                        '[TYPE DE SOCIETE] au capital de [CAPITAL SOCIAL]',
                        modified_text
                    )
                if 'Registre du Commerce' in modified_text:
                    modified_text = re.sub(
                        r'de \[\.+\] sous le numéro \[\.+\]',
                        'de [LOCALITE RCS] sous le numéro [SIREN PRENEUR]',
                        modified_text
                    )
                if 'siège social' in modified_text:
                    modified_text = re.sub(
                        r'à \[\.+\]',
                        'à [ADRESSE DE DOMICILIATION]',
                        modified_text
                    )
                if 'Monsieur' in modified_text or 'Madame' in modified_text:
                    modified_text = re.sub(
                        r'Madame/Monsieur \[\.+\]',
                        '[PRESIDENT DE LA SOCIETE]',
                        modified_text
                    )

        # Remplacements contextuels pour dates, loyers, etc.
        if 'date de prise d\'effet' in modified_text.lower() or 'à compter du' in modified_text:
            modified_text = re.sub(r'\[\.+\]', '[Date de prise d\'effet du bail]', modified_text, count=1)

        if 'durée' in modified_text.lower() and 'années' in modified_text:
            modified_text = re.sub(r'\[\.+\] années', '[Durée initiale du bail] années', modified_text, count=1)

        if 'loyer annuel' in modified_text.lower():
            modified_text = re.sub(r'\[\.+\] euros', '[LOYER ANNUEL HT] euros', modified_text, count=1)

        if 'loyer trimestriel' in modified_text.lower():
            modified_text = re.sub(r'\[\.+\] euros', '[LOYER TRIMESTRIEL] euros', modified_text, count=1)

        if 'dépôt de garantie' in modified_text.lower():
            modified_text = re.sub(r'\[\.+\] euros', '[DEPOT DE GARANTIE] euros', modified_text, count=1)

        # Appliquer les modifications si changement
        if modified_text != original_text:
            para.text = modified_text
            replacements_made += 1

    print(f'   ✅ {replacements_made} paragraphes modifiés')

    # Sauvegarder
    print(f'\n💾 Sauvegarde: {output_docx}')
    doc.save(output_docx)
    print('   ✅ Template créé avec succès!')

    return doc

if __name__ == "__main__":
    input_file = "2024 - Bail type CONVERTED.docx"
    output_file = "Template BAIL avec placeholder V2.docx"

    create_template_with_placeholders(input_file, output_file)

    print('\n' + '=' * 80)
    print('RÉSUMÉ')
    print('=' * 80)
    print(f'✅ Fichier source: {input_file}')
    print(f'✅ Template créé: {output_file}')
    print('\n📋 Prochaines étapes:')
    print('   1. Vérifier le template créé')
    print('   2. Ajuster les placeholders si nécessaire')
    print('   3. Tester la génération de document')
