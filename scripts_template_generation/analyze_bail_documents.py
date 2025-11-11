"""Analyse et comparaison des documents BAIL pour identifier les éléments manquants."""

from docx import Document
import re

def extract_document_structure(doc_path):
    """Extrait la structure d'un document Word."""
    try:
        doc = Document(doc_path)

        structure = {
            "sections": [],
            "articles": [],
            "has_toc": False,
            "has_page_numbers": False,
            "paragraphs": []
        }

        # Analyser les paragraphes
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            structure["paragraphs"].append({
                "index": i,
                "text": text[:200],  # Premiers 200 caractères
                "style": para.style.name if para.style else "Normal"
            })

            # Détecter les articles
            if re.match(r'^ARTICLE\s+\d+', text, re.IGNORECASE):
                structure["articles"].append(text)

            # Détecter table des matières
            if "table des matières" in text.lower() or "sommaire" in text.lower():
                structure["has_toc"] = True

        # Analyser les sections
        for section in doc.sections:
            section_info = {
                "has_header": bool(section.header.paragraphs),
                "has_footer": bool(section.footer.paragraphs),
                "page_height": section.page_height,
                "page_width": section.page_width
            }

            # Vérifier numérotation dans footer
            for para in section.footer.paragraphs:
                if para.text.strip():
                    section_info["footer_text"] = para.text
                    structure["has_page_numbers"] = True

            structure["sections"].append(section_info)

        return structure

    except Exception as e:
        print(f"Erreur lors de l'analyse de {doc_path}: {e}")
        return None

def compare_documents(doc1_path, doc2_path):
    """Compare deux documents et identifie les différences."""
    print("=" * 80)
    print("ANALYSE COMPARATIVE DES DOCUMENTS BAIL")
    print("=" * 80)

    print(f"\n📄 Document 1: {doc1_path}")
    struct1 = extract_document_structure(doc1_path)

    print(f"📄 Document 2: {doc2_path}")
    struct2 = extract_document_structure(doc2_path)

    if not struct1 or not struct2:
        print("❌ Erreur lors de l'analyse")
        return

    # Comparer table des matières
    print("\n" + "=" * 80)
    print("1️⃣  TABLE DES MATIÈRES")
    print("=" * 80)
    print(f"Document 1 (Bail type): {'✅ Présente' if struct1['has_toc'] else '❌ Absente'}")
    print(f"Document 2 (Template): {'✅ Présente' if struct2['has_toc'] else '❌ Absente'}")

    # Comparer numérotation
    print("\n" + "=" * 80)
    print("2️⃣  NUMÉROTATION DES PAGES")
    print("=" * 80)
    print(f"Document 1 (Bail type): {'✅ Présente' if struct1['has_page_numbers'] else '❌ Absente'}")
    print(f"Document 2 (Template): {'✅ Présente' if struct2['has_page_numbers'] else '❌ Absente'}")

    if struct1.get('sections'):
        for i, sec in enumerate(struct1['sections']):
            if 'footer_text' in sec:
                print(f"   Section {i+1} footer: {sec['footer_text']}")

    # Comparer articles
    print("\n" + "=" * 80)
    print("3️⃣  COMPARAISON DES ARTICLES")
    print("=" * 80)
    print(f"Document 1 (Bail type): {len(struct1['articles'])} articles")
    print(f"Document 2 (Template): {len(struct2['articles'])} articles")

    # Trouver articles manquants
    articles1_set = set(struct1['articles'])
    articles2_set = set(struct2['articles'])

    missing_in_template = []
    for art1 in struct1['articles']:
        # Extraire le numéro d'article
        match1 = re.match(r'^ARTICLE\s+(\d+)', art1, re.IGNORECASE)
        if match1:
            num1 = match1.group(1)
            found = False
            for art2 in struct2['articles']:
                match2 = re.match(r'^ARTICLE\s+(\d+)', art2, re.IGNORECASE)
                if match2 and match2.group(1) == num1:
                    found = True
                    break
            if not found:
                missing_in_template.append(art1)

    if missing_in_template:
        print("\n⚠️  Articles présents dans Bail type mais absents du Template:")
        for art in missing_in_template:
            print(f"   - {art}")
    else:
        print("\n✅ Tous les articles du Bail type sont présents dans le Template")

    # Afficher tous les articles du Bail type
    print("\n" + "=" * 80)
    print("4️⃣  LISTE COMPLÈTE DES ARTICLES (Bail type)")
    print("=" * 80)
    for i, art in enumerate(struct1['articles'], 1):
        print(f"{i:2d}. {art}")

    # Afficher le contenu détaillé du début du document
    print("\n" + "=" * 80)
    print("5️⃣  STRUCTURE DU DÉBUT DU DOCUMENT (Bail type)")
    print("=" * 80)
    for i, para in enumerate(struct1['paragraphs'][:20]):  # 20 premiers paragraphes
        if para['text']:
            print(f"[{i}] Style: {para['style']}")
            print(f"    {para['text'][:150]}")
            print()

if __name__ == "__main__":
    doc1 = "2024 - Bail type.doc"
    doc2 = "Template BAIL avec placeholder.docx"

    compare_documents(doc1, doc2)
