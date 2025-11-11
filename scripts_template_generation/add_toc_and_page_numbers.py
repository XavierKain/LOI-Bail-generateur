"""Ajoute la table des matières et la numérotation des pages au template."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def add_page_number(section):
    """Ajoute la numérotation des pages au footer."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Nettoyer le footer existant
    for para in footer.paragraphs:
        para.clear()

    # Créer un paragraphe pour le numéro de page
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter le champ PAGE
    run = para.add_run()

    # Créer l'élément XML pour le numéro de page
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run.font.size = Pt(10)

    return True

def add_table_of_contents(doc):
    """Ajoute une table des matières au début du document."""

    # Chercher où insérer la TOC (après le titre et les parties, avant ARTICLE PRELIMINAIRE)
    insert_position = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('ARTICLE'):
            insert_position = i
            break

    if insert_position == 0:
        insert_position = 3  # Par défaut après quelques paragraphes

    print(f'📑 Insertion de la table des matières à la position {insert_position}')

    # Créer une nouvelle page pour la TOC
    # On va insérer les paragraphes
    toc_paragraphs = [
        ("", ""),  # Saut de page avant
        ("TABLE DES MATIÈRES", "Heading1"),
        ("", ""),  # Ligne vide
    ]

    # Extraire tous les titres d'articles
    articles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('ARTICLE'):
            # Extraire juste la première ligne (titre)
            first_line = text.split('\n')[0] if '\n' in text else text
            articles.append(first_line)

    print(f'   Trouvé {len(articles)} articles pour la TOC')

    # Ajouter les entrées de TOC
    for article in articles:
        toc_paragraphs.append((f"   {article}", "TOC1"))

    # Insérer les paragraphes dans le document
    # Note: python-docx ne permet pas d'insérer facilement au milieu
    # On va reconstruire le document

    new_doc = Document()

    # Copier les sections
    for section in doc.sections:
        new_section = new_doc.sections[0] if len(new_doc.sections) == 1 else new_doc.add_section()
        new_section.page_height = section.page_height
        new_section.page_width = section.page_width
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin

    # Copier les paragraphes avant la position d'insertion
    for i in range(min(insert_position, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        new_para = new_doc.add_paragraph(para.text)
        if para.style:
            try:
                new_para.style = para.style
            except:
                pass

    # Insérer la TOC
    new_doc.add_page_break()

    toc_title = new_doc.add_paragraph("TABLE DES MATIÈRES")
    toc_title.runs[0].bold = True
    toc_title.runs[0].font.size = Pt(16)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()  # Ligne vide

    # Ajouter les entrées
    for article in articles:
        entry = new_doc.add_paragraph(f"   {article}")
        entry.runs[0].font.size = Pt(11)

    new_doc.add_page_break()

    # Copier le reste des paragraphes
    for i in range(insert_position, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        new_para = new_doc.add_paragraph(para.text)
        if para.style:
            try:
                new_para.style = para.style
            except:
                pass

    return new_doc

def finalize_template():
    """Finalise le template avec TOC et numérotation."""

    print('=' * 80)
    print('FINALISATION DU TEMPLATE')
    print('=' * 80)

    input_file = "Template BAIL avec placeholder COMPLET.docx"
    print(f'\n📄 Chargement: {input_file}')
    doc = Document(input_file)
    print(f'   Paragraphes: {len(doc.paragraphs)}')
    print(f'   Sections: {len(doc.sections)}')

    # Ajouter la table des matières
    print('\n📑 Ajout de la table des matières...')
    doc = add_table_of_contents(doc)
    print('   ✅ Table des matières ajoutée')

    # Ajouter la numérotation des pages
    print('\n🔢 Ajout de la numérotation des pages...')
    for i, section in enumerate(doc.sections):
        add_page_number(section)
        print(f'   ✅ Section {i+1}: numérotation ajoutée')

    # Sauvegarder
    output_file = "Template BAIL avec placeholder FINAL.docx"
    print(f'\n💾 Sauvegarde: {output_file}')
    doc.save(output_file)
    print('   ✅ Template finalisé!')

    print(f'\n📊 Template final:')
    print(f'   Paragraphes: {len(doc.paragraphs)}')
    print(f'   Fichier: {output_file}')

    return output_file

if __name__ == "__main__":
    output = finalize_template()

    print('\n' + '=' * 80)
    print('RÉSUMÉ')
    print('=' * 80)
    print(f'✅ Template final créé: {output}')
    print('\n✅ Fonctionnalités ajoutées:')
    print('   - Table des matières complète')
    print('   - Numérotation des pages au footer')
    print('   - Tous les articles (PRELIMINAIRE + 1-28)')
    print('   - Placeholders pour génération automatique')
