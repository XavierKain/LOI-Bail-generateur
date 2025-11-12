"""Test rapide du nouveau template BAIL complet."""

from docx import Document
import re

def test_template():
    """Teste le template FINAL."""

    print('=' * 80)
    print('TEST DU TEMPLATE BAIL FINAL')
    print('=' * 80)

    template_path = "Template BAIL avec placeholder AUTOMATISE.docx"

    print(f'\n📄 Chargement: {template_path}')
    doc = Document(template_path)

    print(f'\n📊 Structure du document:')
    print(f'   Paragraphes: {len(doc.paragraphs)}')
    print(f'   Sections: {len(doc.sections)}')

    # Vérifier la table des matières
    print('\n📑 Vérification de la table des matières:')
    toc_found = False
    for para in doc.paragraphs[:50]:  # Chercher dans les 50 premiers paragraphes
        if 'TABLE DES MATIÈRES' in para.text.upper() or 'TABLE DES MATIERES' in para.text.upper():
            toc_found = True
            print('   ✅ Table des matières trouvée')
            break

    if not toc_found:
        print('   ⚠️  Table des matières non trouvée')

    # Vérifier les numéros de page
    print('\n🔢 Vérification de la numérotation des pages:')
    has_page_numbers = False
    for section in doc.sections:
        if section.footer.paragraphs:
            for para in section.footer.paragraphs:
                if para.text.strip() or len(para._element.xpath('.//w:fldChar')) > 0:
                    has_page_numbers = True
                    print('   ✅ Numérotation des pages présente')
                    break
        if has_page_numbers:
            break

    if not has_page_numbers:
        print('   ⚠️  Numérotation des pages non trouvée')

    # Compter les articles
    print('\n📋 Vérification des articles:')
    articles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('ARTICLE'):
            # Extraire juste la première ligne
            first_line = text.split('\n')[0] if '\n' in text else text
            match = re.match(r'^ARTICLE\s+(PRELIMINAIRE|\d+)', first_line)
            if match:
                articles.append(match.group(1))

    print(f'   Articles trouvés: {len(articles)}')

    # Vérifier la présence de tous les articles
    expected = ['PRELIMINAIRE'] + [str(i) for i in range(1, 29)]
    found_nums = set(articles)
    expected_nums = set(expected)

    missing = expected_nums - found_nums
    if missing:
        print(f'   ⚠️  Articles manquants: {sorted(missing)}')
    else:
        print('   ✅ Tous les articles présents (PRELIMINAIRE + 1-28)')

    extra = found_nums - expected_nums
    if extra:
        print(f'   ⚠️  Articles supplémentaires: {sorted(extra)}')

    # Vérifier les placeholders clés
    print('\n🔖 Vérification des placeholders clés:')
    full_text = '\n'.join([p.text for p in doc.paragraphs])

    key_placeholders = [
        '[NOM DU BAILLEUR]',
        '[SIREN BAILLEUR]',
        '[NOM DU PRENEUR]',
        '[SIREN PRENEUR]',
        '[CAPITAL SOCIAL]',
        '[LOYER ANNUEL HT]',
        '[Date de prise d\'effet du bail]',
        '[Durée initiale du bail]',
        '[PRESIDENT DE LA SOCIETE]',
    ]

    found_placeholders = 0
    for placeholder in key_placeholders:
        if placeholder in full_text:
            found_placeholders += 1
            print(f'   ✅ {placeholder}')
        else:
            print(f'   ❌ {placeholder} NON TROUVÉ')

    print(f'\n   Total: {found_placeholders}/{len(key_placeholders)} placeholders trouvés')

    # Résumé
    print('\n' + '=' * 80)
    print('RÉSUMÉ DU TEST')
    print('=' * 80)

    all_checks = [
        ('Table des matières', toc_found),
        ('Numérotation pages', has_page_numbers),
        ('Tous les articles', len(missing) == 0),
        ('Placeholders clés', found_placeholders >= 7),
    ]

    passed = sum(1 for _, check in all_checks if check)
    total = len(all_checks)

    print(f'\nTests réussis: {passed}/{total}')
    for name, check in all_checks:
        status = '✅' if check else '❌'
        print(f'   {status} {name}')

    if passed == total:
        print('\n🎉 TEMPLATE VALIDE - Prêt pour la génération!')
    else:
        print('\n⚠️  TEMPLATE INCOMPLET - Ajustements nécessaires')

    return passed == total

if __name__ == "__main__":
    success = test_template()
    exit(0 if success else 1)
